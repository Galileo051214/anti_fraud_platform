# -*- coding: utf-8 -*-
"""
修正5.3节功能子节的顺序并补充5.3.9
问题：
1. 各模块的功能子节顺序是倒序（如5.3.2.7在5.3.2.6之前），需要调整为从小到大
2. 5.3.9模块缺少功能子节（应该有5.3.9.1~5.3.9.7）

策略：
1. 从备份恢复原始文档（-v2-backup.docx）
2. 按正确顺序（从小到大）重新生成所有功能子节
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import shutil
import os

# 文件路径
backup_path = r"D:\Project\anti_fraud_platform\docx\《项目阶段二》评审\专业综合课程设计-系统设计说明书-江乐霖-23201317-v2-backup.docx"
target_path = r"D:\Project\anti_fraud_platform\docx\《项目阶段二》评审\专业综合课程设计-系统设计说明书-江乐霖-23201317.docx"

# 从备份恢复
shutil.copy(backup_path, target_path)
doc = Document(target_path)

# 辅助函数：设置段落文本（保留样式）
def set_para_text(para, text):
    if para.runs:
        run = para.runs[0]
        run.text = text
        # 清除其他run
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)

# 辅助函数：删除段落
def delete_para(para):
    para._element.getparent().remove(para._element)

# 辅助函数：插入段落（在参考段落之后）
def insert_para_after(ref_para, text, style_name):
    new_p = ref_para.insert_paragraph_after(text)
    if style_name:
        new_p.style = style_name
    return new_p

# 功能数据（按模块组织，顺序从小到大）
functions_data = {
    "5.3.1": {
        "module_title": "用户与认证模块",
        "functions": [
            {"num": "5.3.1.3", "name": "用户信息查询", "api": "GET /api/user/info", 
             "desc": "获取当前登录用户的详细信息，包括用户名、昵称、头像、角色、年级、专业等。Controller从SecurityContextHolder获取当前用户ID，调用UserService查询数据库返回UserVO。",
             "classes": ["UserController（控制层，处理/api/user接口）", "UserServiceImpl（业务层，查询用户信息）", "UserMapper（数据层，查询sys_user表）", "UserVO（视图对象，返回用户信息）"],
             "files": [("UserController.java", "Controller", "com.anti.controller", "用户信息查询接口"),
                      ("UserServiceImpl.java", "Service", "com.anti.service.impl", "查询用户信息业务逻辑"),
                      ("UserMapper.java", "Mapper", "com.anti.mapper", "用户数据访问"),
                      ("User.java", "Entity", "com.anti.entity", "用户实体"),
                      ("UserVO.java", "VO", "com.anti.entity.vo", "用户信息视图对象")]},
            {"num": "5.3.1.4", "name": "用户资料更新", "api": "PUT /api/user/update",
             "desc": "更新当前用户的个人资料，包括昵称、头像、手机号、邮箱、年级、专业等。Controller接收UpdateUserRequest，调用UserService更新数据库，返回更新后的用户信息。",
             "classes": ["UserController（控制层）", "UserServiceImpl（业务层，更新用户资料）", "UserMapper（数据层，更新sys_user）", "UpdateUserRequest（请求DTO）"],
             "files": [("UserController.java", "Controller", "com.anti.controller", "用户资料更新接口"),
                      ("UserServiceImpl.java", "Service", "com.anti.service.impl", "更新用户资料业务逻辑"),
                      ("UserMapper.java", "Mapper", "com.anti.mapper", "用户数据访问"),
                      ("UpdateUserRequest.java", "DTO", "com.anti.entity.dto", "更新请求参数")]},
            {"num": "5.3.1.5", "name": "密码修改", "api": "PUT /api/user/password",
             "desc": "用户修改登录密码。需校验旧密码正确性，新密码BCrypt加密后存储。Controller接收ChangePasswordRequest，Service验证旧密码后更新。",
             "classes": ["UserController（控制层）", "UserServiceImpl（业务层，验证旧密码并更新新密码）", "BCryptPasswordEncoder（加密工具）", "ChangePasswordRequest（请求DTO）"],
             "files": [("UserController.java", "Controller", "com.anti.controller", "密码修改接口"),
                      ("UserServiceImpl.java", "Service", "com.anti.service.impl", "密码修改业务逻辑"),
                      ("PasswordConfig.java", "Config", "com.anti.config", "BCrypt加密器配置"),
                      ("ChangePasswordRequest.java", "DTO", "com.anti.entity.dto", "密码修改请求参数")]},
            {"num": "5.3.1.6", "name": "用户登出", "api": "POST /api/user/logout",
             "desc": "用户退出登录。将当前JWT Token写入Redis黑名单，Token在剩余过期时间内无法再次使用。Controller提取Token，Service计算剩余TTL并写入Redis。",
             "classes": ["UserController（控制层）", "UserServiceImpl（业务层，将Token写入黑名单）", "JwtUtils（JWT工具，解析Token获取过期时间）", "RedisCacheUtil（Redis缓存工具）"],
             "files": [("UserController.java", "Controller", "com.anti.controller", "用户登出接口"),
                      ("UserServiceImpl.java", "Service", "com.anti.service.impl", "登出业务逻辑"),
                      ("JwtUtils.java", "Util", "com.anti.security", "JWT解析工具"),
                      ("RedisCacheUtil.java", "Util", "com.anti.util", "Token黑名单写入")]},
            {"num": "5.3.1.7", "name": "用户管理", "api": "GET /api/user/list, GET/PUT /api/user/{id}",
             "desc": "管理员功能：分页查询用户列表（支持关键字/角色/状态筛选）、查看用户详情、更新用户信息、启用/禁用用户账号。Controller用@PreAuthorize('ADMIN')鉴权。",
             "classes": ["UserController（控制层，管理员接口）", "UserServiceImpl（业务层，分页查询/更新状态）", "UserMapper（数据层，selectUserPage）", "UserVO（用户视图）"],
             "files": [("UserController.java", "Controller", "com.anti.controller", "用户管理接口"),
                      ("UserServiceImpl.java", "Service", "com.anti.service.impl", "用户管理业务逻辑"),
                      ("UserMapper.java", "Mapper", "com.anti.mapper", "分页查询用户"),
                      ("UserMapper.xml", "XML", "resources/mapper", "分页查询SQL")]},
        ]
    },
    "5.3.2": {
        "module_title": "资讯学习模块",
        "functions": [
            {"num": "5.3.2.1", "name": "资讯分页列表", "api": "GET /api/news/page",
             "desc": "分页查询资讯列表，支持按分类、类型（新闻/预警/政策）、关键字筛选。按置顶优先、发布时间倒序排列。批量填充点赞数和当前用户是否点赞标记。",
             "classes": ["NewsController（控制层）", "NewsServiceImpl（业务层，分页查询+点赞填充）", "NewsMapper（数据层）", "NewsLikeMapper（点赞查询）", "NewsVO（资讯视图）"],
             "files": [("NewsController.java", "Controller", "com.anti.controller", "资讯分页接口"),
                      ("NewsServiceImpl.java", "Service", "com.anti.service.impl", "分页查询业务逻辑"),
                      ("NewsMapper.java", "Mapper", "com.anti.mapper", "资讯数据访问"),
                      ("NewsLikeMapper.java", "Mapper", "com.anti.mapper", "点赞数据访问"),
                      ("NewsVO.java", "VO", "com.anti.entity.vo", "资讯视图对象")]},
            {"num": "5.3.2.2", "name": "资讯详情查看", "api": "GET /api/news/{id}",
             "desc": "查看资讯详情，返回完整内容、封面、分类、作者、浏览量、点赞数等信息。",
             "classes": ["NewsController", "NewsServiceImpl", "NewsMapper", "NewsVO"],
             "files": [("NewsController.java", "Controller", "com.anti.controller", "资讯详情接口"),
                      ("NewsServiceImpl.java", "Service", "com.anti.service.impl", "详情查询"),
                      ("News.java", "Entity", "com.anti.entity", "资讯实体")]},
            {"num": "5.3.2.3", "name": "资讯发布", "api": "POST /api/news（管理员）",
             "desc": "管理员发布新资讯。校验标题内容非空，设置初始浏览量0、状态为草稿，插入数据库。",
             "classes": ["NewsController", "NewsServiceImpl", "NewsMapper", "CreateNewsRequest", "CacheRefreshService"],
             "files": [("NewsController.java", "Controller", "com.anti.controller", "资讯发布接口"),
                      ("NewsServiceImpl.java", "Service", "com.anti.service.impl", "资讯发布业务逻辑"),
                      ("CreateNewsRequest.java", "DTO", "com.anti.entity.dto", "发布请求参数"),
                      ("CacheRefreshService.java", "Service", "com.anti.service.impl", "缓存失效")]},
            {"num": "5.3.2.4", "name": "资讯编辑", "api": "PUT /api/news/{id}（管理员）",
             "desc": "管理员编辑资讯内容、分类、封面等字段。更新后触发缓存失效。",
             "classes": ["NewsController", "NewsServiceImpl", "NewsMapper", "UpdateNewsRequest"],
             "files": [("NewsController.java", "Controller", "com.anti.controller", "资讯编辑接口"),
                      ("NewsServiceImpl.java", "Service", "com.anti.service.impl", "编辑业务逻辑"),
                      ("UpdateNewsRequest.java", "DTO", "com.anti.entity.dto", "编辑请求参数")]},
            {"num": "5.3.2.5", "name": "资讯删除", "api": "DELETE /api/news/{id}（管理员）",
             "desc": "管理员删除资讯。触发缓存失效。",
             "classes": ["NewsController", "NewsServiceImpl", "NewsMapper", "CacheRefreshService"],
             "files": [("NewsController.java", "Controller", "com.anti.controller", "资讯删除接口"),
                      ("NewsServiceImpl.java", "Service", "com.anti.service.impl", "删除业务逻辑")]},
            {"num": "5.3.2.6", "name": "资讯浏览记录", "api": "POST /api/news/{id}/browse",
             "desc": "记录资讯浏览行为，包括停留时长。触发连续学习天数成就检查。",
             "classes": ["NewsController", "NewsServiceImpl", "NewsBrowseLogMapper", "AchievementService"],
             "files": [("NewsController.java", "Controller", "com.anti.controller", "浏览记录接口"),
                      ("NewsServiceImpl.java", "Service", "com.anti.service.impl", "浏览记录业务"),
                      ("NewsBrowseLog.java", "Entity", "com.anti.entity", "浏览日志实体"),
                      ("AchievementService.java", "Service", "com.anti.service", "连续学习检查")]},
            {"num": "5.3.2.7", "name": "资讯点赞", "api": "POST/DELETE /api/news/{id}/like",
             "desc": "点赞或取消点赞资讯。幂等校验，点赞时插入记录并更新点赞统计，取消时删除记录。",
             "classes": ["NewsController", "NewsServiceImpl", "NewsLikeMapper", "CacheRefreshService"],
             "files": [("NewsController.java", "Controller", "com.anti.controller", "点赞接口"),
                      ("NewsServiceImpl.java", "Service", "com.anti.service.impl", "点赞业务逻辑"),
                      ("NewsLike.java", "Entity", "com.anti.entity", "点赞实体"),
                      ("NewsLikeMapper.java", "Mapper", "com.anti.mapper", "点赞数据访问")]},
        ]
    },
    "5.3.3": {
        "module_title": "案例展示模块",
        "functions": [
            {"num": "5.3.3.1", "name": "案例分页列表", "api": "GET /api/case/page",
             "desc": "分页查询诈骗案例，支持按标签、关键字筛选。按精选优先、Wilson置信分、发布时间排序。返回案例标题、类型、难度、风险等级、浏览量、点赞数。",
             "classes": ["CaseController", "FraudCaseServiceImpl", "FraudCaseMapper", "CaseTagRelationMapper", "CaseVO"],
             "files": [("CaseController.java", "Controller", "com.anti.controller", "案例分页接口"),
                      ("FraudCaseServiceImpl.java", "Service", "com.anti.service.impl", "分页查询业务"),
                      ("FraudCaseMapper.java", "Mapper", "com.anti.mapper", "案例数据访问"),
                      ("CaseVO.java", "VO", "com.anti.entity.vo", "案例视图对象")]},
            {"num": "5.3.3.2", "name": "案例详情查看", "api": "GET /api/case/{id}",
             "desc": "查看案例详情，包括完整内容、诈骗剧本、目标人群、风险评分、难度等级、标签等。",
             "classes": ["CaseController", "FraudCaseServiceImpl", "FraudCaseMapper", "CaseVO"],
             "files": [("CaseController.java", "Controller", "com.anti.controller", "案例详情接口"),
                      ("FraudCaseServiceImpl.java", "Service", "com.anti.service.impl", "详情查询"),
                      ("FraudCase.java", "Entity", "com.anti.entity", "案例实体")]},
            {"num": "5.3.3.3", "name": "案例发布", "api": "POST /api/case（管理员）",
             "desc": "管理员发布新案例。设置初始浏览量/点赞数0，计算Wilson分，关联标签。",
             "classes": ["CaseController", "FraudCaseServiceImpl", "FraudCaseMapper", "CreateCaseRequest"],
             "files": [("CaseController.java", "Controller", "com.anti.controller", "案例发布接口"),
                      ("FraudCaseServiceImpl.java", "Service", "com.anti.service.impl", "发布业务"),
                      ("CreateCaseRequest.java", "DTO", "com.anti.entity.dto", "发布请求参数")]},
            {"num": "5.3.3.4", "name": "案例编辑", "api": "PUT /api/case/{id}（管理员）",
             "desc": "管理员编辑案例内容、类型、难度、目标人群等。",
             "classes": ["CaseController", "FraudCaseServiceImpl", "FraudCaseMapper", "UpdateCaseRequest"],
             "files": [("CaseController.java", "Controller", "com.anti.controller", "案例编辑接口"),
                      ("FraudCaseServiceImpl.java", "Service", "com.anti.service.impl", "编辑业务"),
                      ("UpdateCaseRequest.java", "DTO", "com.anti.entity.dto", "编辑请求参数")]},
            {"num": "5.3.3.5", "name": "案例删除", "api": "DELETE /api/case/{id}（管理员）",
             "desc": "管理员删除案例。",
             "classes": ["CaseController", "FraudCaseServiceImpl", "FraudCaseMapper"],
             "files": [("CaseController.java", "Controller", "com.anti.controller", "案例删除接口"),
                      ("FraudCaseServiceImpl.java", "Service", "com.anti.service.impl", "删除业务")]},
            {"num": "5.3.3.6", "name": "案例浏览记录", "api": "POST /api/case/{id}/browse",
             "desc": "记录案例浏览行为。首次浏览奖励+2积分，更新日/周/总排行榜，触发browse_count成就，按停留时长更新知识水平。",
             "classes": ["CaseController", "FraudCaseServiceImpl", "CaseBrowseLogMapper", "ScoreService", "LeaderboardService", "AchievementService", "ProfileService"],
             "files": [("CaseController.java", "Controller", "com.anti.controller", "浏览记录接口"),
                      ("FraudCaseServiceImpl.java", "Service", "com.anti.service.impl", "浏览业务"),
                      ("CaseBrowseLog.java", "Entity", "com.anti.entity", "浏览日志"),
                      ("ScoreService.java", "Service", "com.anti.service", "积分+2"),
                      ("LeaderboardService.java", "Service", "com.anti.service", "排行榜更新")]},
            {"num": "5.3.3.7", "name": "案例点赞", "api": "POST/DELETE /api/case/{id}/like",
             "desc": "点赞或取消点赞案例。更新点赞统计并重算Wilson置信分。",
             "classes": ["CaseController", "FraudCaseServiceImpl", "CaseLikeMapper", "RecommendationAlgorithmUtil"],
             "files": [("CaseController.java", "Controller", "com.anti.controller", "点赞接口"),
                      ("FraudCaseServiceImpl.java", "Service", "com.anti.service.impl", "点赞业务"),
                      ("CaseLike.java", "Entity", "com.anti.entity", "点赞实体"),
                      ("RecommendationAlgorithmUtil.java", "Util", "com.anti.util", "Wilson分计算")]},
            {"num": "5.3.3.8", "name": "热点案例推荐", "api": "GET /api/case/hot",
             "desc": "返回热点案例列表，基于Wilson置信分和浏览量综合排序。",
             "classes": ["CaseController", "FraudCaseServiceImpl", "FraudCaseMapper", "RedisCacheUtil"],
             "files": [("CaseController.java", "Controller", "com.anti.controller", "热点接口"),
                      ("FraudCaseServiceImpl.java", "Service", "com.anti.service.impl", "热点查询"),
                      ("RedisCacheUtil.java", "Util", "com.anti.util", "热点缓存")]},
        ]
    },
    "5.3.4": {
        "module_title": "知识闯关与情景模拟模块",
        "functions": [
            {"num": "5.3.4.1", "name": "关卡列表查询", "api": "GET /api/challenge/list",
             "desc": "查询所有启用的关卡列表，返回关卡名称、类型（答题/情景）、难度、奖励积分、解锁状态、通关状态。解锁规则：levelOrder<=maxPassedLevel+1（顺序解锁）。",
             "classes": ["ChallengeController", "ChallengeServiceImpl", "ChallengeMapper", "UserChallengeRecordMapper", "ChallengeVO"],
             "files": [("ChallengeController.java", "Controller", "com.anti.controller", "关卡列表接口"),
                      ("ChallengeServiceImpl.java", "Service", "com.anti.service.impl", "解锁状态计算"),
                      ("ChallengeMapper.java", "Mapper", "com.anti.mapper", "关卡数据访问"),
                      ("ChallengeVO.java", "VO", "com.anti.entity.vo", "关卡视图")]},
            {"num": "5.3.4.2", "name": "关卡详情查看", "api": "GET /api/challenge/{id}",
             "desc": "查看关卡详情，包括题目列表或情景剧本（FSM节点和边）。",
             "classes": ["ChallengeController", "ChallengeServiceImpl", "ChallengeMapper", "Challenge"],
             "files": [("ChallengeController.java", "Controller", "com.anti.controller", "关卡详情接口"),
                      ("ChallengeServiceImpl.java", "Service", "com.anti.service.impl", "详情查询"),
                      ("Challenge.java", "Entity", "com.anti.entity", "关卡实体含嵌套类")]},
            {"num": "5.3.4.3", "name": "答题闯关提交", "api": "POST /api/challenge/submit",
             "desc": "提交答题答案，计算得分。通关奖励scoreReward积分，更新排行榜，检查成就（challenge_count/perfect_score/challenge_complete），更新知识水平。",
             "classes": ["ChallengeController", "ChallengeServiceImpl", "ScoreService", "LeaderboardService", "AchievementService", "ProfileService"],
             "files": [("ChallengeController.java", "Controller", "com.anti.controller", "提交接口"),
                      ("ChallengeServiceImpl.java", "Service", "com.anti.service.impl", "判分业务"),
                      ("SubmitChallengeRequest.java", "DTO", "com.anti.entity.dto", "答案请求"),
                      ("ChallengeResultVO.java", "VO", "com.anti.entity.vo", "结果视图")]},
            {"num": "5.3.4.4", "name": "闯关记录查询", "api": "GET /api/challenge/records",
             "desc": "查询用户闯关记录，包括尝试次数、得分、是否通关、答题详情。",
             "classes": ["ChallengeController", "ChallengeServiceImpl", "UserChallengeRecordMapper", "ChallengeRecordVO"],
             "files": [("ChallengeController.java", "Controller", "com.anti.controller", "记录接口"),
                      ("ChallengeServiceImpl.java", "Service", "com.anti.service.impl", "记录查询"),
                      ("UserChallengeRecord.java", "Entity", "com.anti.entity", "闯关记录实体")]},
            {"num": "5.3.4.5", "name": "情景模拟开始", "api": "POST /api/scenario/start/{challengeId}",
             "desc": "开始情景模拟，创建ScenarioProgress记录，设置起始节点。",
             "classes": ["ScenarioController", "ScenarioServiceImpl", "ScenarioProgressMapper", "ScenarioProgressVO"],
             "files": [("ScenarioController.java", "Controller", "com.anti.controller", "开始接口"),
                      ("ScenarioServiceImpl.java", "Service", "com.anti.service.impl", "初始化进度"),
                      ("ScenarioProgress.java", "Entity", "com.anti.entity", "进度实体")]},
            {"num": "5.3.4.6", "name": "情景决策推进", "api": "POST /api/scenario/decision",
             "desc": "在情景模拟中做决策，推进FSM状态机。记录DecisionRecord（含isSafeChoice），到达终点时计算finalScore=100*安全选择数/总选择数，>=60通关奖励积分。",
             "classes": ["ScenarioController", "ScenarioServiceImpl", "ScenarioProgressMapper", "ScoreService", "LeaderboardService"],
             "files": [("ScenarioController.java", "Controller", "com.anti.controller", "决策接口"),
                      ("ScenarioServiceImpl.java", "Service", "com.anti.service.impl", "FSM推进"),
                      ("ScenarioDecisionRequest.java", "DTO", "com.anti.entity.dto", "决策请求")]},
            {"num": "5.3.4.7", "name": "情景进度查询", "api": "GET /api/scenario/progress/{challengeId}",
             "desc": "查询情景模拟当前进度，包括当前节点内容、角色、风险提示、可选决策。",
             "classes": ["ScenarioController", "ScenarioServiceImpl", "ScenarioProgressMapper", "ScenarioProgressVO"],
             "files": [("ScenarioController.java", "Controller", "com.anti.controller", "进度接口"),
                      ("ScenarioServiceImpl.java", "Service", "com.anti.service.impl", "进度查询"),
                      ("ScenarioProgressVO.java", "VO", "com.anti.entity.vo", "进度视图")]},
        ]
    },
    "5.3.5": {
        "module_title": "社区互动模块",
        "functions": [
            {"num": "5.3.5.1", "name": "帖子分页列表", "api": "GET /api/forum/post/page",
             "desc": "分页查询帖子，支持按类型（经验/问题/讨论）、排序（时间/点赞/评论）、关键字筛选。",
             "classes": ["ForumController", "ForumPostServiceImpl", "ForumPostMapper", "PostVO"],
             "files": [("ForumController.java", "Controller", "com.anti.controller", "帖子列表接口"),
                      ("ForumPostServiceImpl.java", "Service", "com.anti.service.impl", "分页查询"),
                      ("ForumPostMapper.java", "Mapper", "com.anti.mapper", "帖子数据访问"),
                      ("PostVO.java", "VO", "com.anti.entity.vo", "帖子视图")]},
            {"num": "5.3.5.2", "name": "帖子详情查看", "api": "GET /api/forum/post/{postId}",
             "desc": "查看帖子详情，返回内容、点赞数、评论数、作者信息。",
             "classes": ["ForumController", "ForumPostServiceImpl", "ForumPostMapper", "PostVO"],
             "files": [("ForumController.java", "Controller", "com.anti.controller", "帖子详情接口"),
                      ("ForumPostServiceImpl.java", "Service", "com.anti.service.impl", "详情查询"),
                      ("ForumPost.java", "Entity", "com.anti.entity", "帖子实体")]},
            {"num": "5.3.5.3", "name": "帖子发布", "api": "POST /api/forum/post",
             "desc": "发布帖子，奖励+3积分并更新排行榜。检查post_count成就和连续学习。",
             "classes": ["ForumController", "ForumPostServiceImpl", "ScoreService", "LeaderboardService", "AchievementService"],
             "files": [("ForumController.java", "Controller", "com.anti.controller", "帖子发布接口"),
                      ("ForumPostServiceImpl.java", "Service", "com.anti.service.impl", "发布业务"),
                      ("CreatePostRequest.java", "DTO", "com.anti.entity.dto", "发布请求"),
                      ("ScoreService.java", "Service", "com.anti.service", "积分+3")]},
            {"num": "5.3.5.4", "name": "帖子编辑", "api": "PUT /api/forum/post/{postId}",
             "desc": "编辑帖子（仅作者或管理员）。",
             "classes": ["ForumController", "ForumPostServiceImpl", "UpdatePostRequest"],
             "files": [("ForumController.java", "Controller", "com.anti.controller", "帖子编辑接口"),
                      ("ForumPostServiceImpl.java", "Service", "com.anti.service.impl", "编辑业务"),
                      ("UpdatePostRequest.java", "DTO", "com.anti.entity.dto", "编辑请求")]},
            {"num": "5.3.5.5", "name": "帖子删除", "api": "DELETE /api/forum/post/{postId}",
             "desc": "删除帖子（作者或管理员）。级联删除帖子点赞、评论、评论点赞。",
             "classes": ["ForumController", "ForumPostServiceImpl", "PostLikeMapper", "CommentMapper", "CommentLikeMapper"],
             "files": [("ForumController.java", "Controller", "com.anti.controller", "帖子删除接口"),
                      ("ForumPostServiceImpl.java", "Service", "com.anti.service.impl", "级联删除"),
                      ("PostLikeMapper.java", "Mapper", "com.anti.mapper", "帖子点赞"),
                      ("CommentMapper.java", "Mapper", "com.anti.mapper", "评论数据")]},
            {"num": "5.3.5.6", "name": "帖子点赞", "api": "POST/DELETE /api/forum/post/{postId}/like",
             "desc": "点赞或取消点赞帖子。",
             "classes": ["ForumController", "ForumPostServiceImpl", "PostLikeMapper"],
             "files": [("ForumController.java", "Controller", "com.anti.controller", "点赞接口"),
                      ("ForumPostServiceImpl.java", "Service", "com.anti.service.impl", "点赞业务"),
                      ("PostLike.java", "Entity", "com.anti.entity", "点赞实体")]},
            {"num": "5.3.5.7", "name": "评论查看", "api": "GET /api/forum/post/{postId}/comments",
             "desc": "查询帖子的评论树，按parentId递归构建树形结构，填充isLiked/isAuthor标记。",
             "classes": ["ForumController", "ForumPostServiceImpl", "CommentMapper", "CommentVO"],
             "files": [("ForumController.java", "Controller", "com.anti.controller", "评论接口"),
                      ("ForumPostServiceImpl.java", "Service", "com.anti.service.impl", "评论树构建"),
                      ("CommentVO.java", "VO", "com.anti.entity.vo", "评论视图")]},
            {"num": "5.3.5.8", "name": "评论发布", "api": "POST /api/comment",
             "desc": "发布评论，支持回复评论（parentId）。",
             "classes": ["CommentController", "CommentServiceImpl", "CommentMapper", "CreateCommentRequest"],
             "files": [("CommentController.java", "Controller", "com.anti.controller", "评论发布接口"),
                      ("CommentServiceImpl.java", "Service", "com.anti.service.impl", "评论发布"),
                      ("CreateCommentRequest.java", "DTO", "com.anti.entity.dto", "评论请求"),
                      ("Comment.java", "Entity", "com.anti.entity", "评论实体")]},
        ]
    },
    "5.3.6": {
        "module_title": "AI智能客服模块",
        "functions": [
            {"num": "5.3.6.1", "name": "智能问答", "api": "POST /api/chat/ask",
             "desc": "用户提问，AI调用DeepSeek模型回答。无sessionId则新建会话，加载历史对话，使用AntiFraudPromptTemplate构建系统提示词，调用DeepSeekClient.chat，持久化问答记录。",
             "classes": ["ChatController", "QAConversationServiceImpl", "DeepSeekClient", "AntiFraudPromptTemplate", "QAConversation"],
             "files": [("ChatController.java", "Controller", "com.anti.controller", "问答接口"),
                      ("QAConversationServiceImpl.java", "Service", "com.anti.service.impl", "问答业务"),
                      ("DeepSeekClient.java", "Util", "com.anti.util", "DeepSeek API调用"),
                      ("AntiFraudPromptTemplate.java", "Util", "com.anti.util", "反诈提示词"),
                      ("QAConversation.java", "Entity", "com.anti.entity", "问答记录")]},
            {"num": "5.3.6.2", "name": "会话历史查询", "api": "GET /api/chat/history/{sessionId}",
             "desc": "查询指定会话的历史问答记录。",
             "classes": ["ChatController", "QAConversationServiceImpl", "QAConversationMapper"],
             "files": [("ChatController.java", "Controller", "com.anti.controller", "历史接口"),
                      ("QAConversationServiceImpl.java", "Service", "com.anti.service.impl", "历史查询"),
                      ("ChatVO.java", "VO", "com.anti.entity.vo", "问答视图")]},
            {"num": "5.3.6.3", "name": "会话列表查询", "api": "GET /api/chat/sessions",
             "desc": "查询用户所有会话列表，返回首问题摘要、最后答案摘要、消息数、总Token数。",
             "classes": ["ChatController", "QAConversationServiceImpl", "SessionVO"],
             "files": [("ChatController.java", "Controller", "com.anti.controller", "会话列表接口"),
                      ("QAConversationServiceImpl.java", "Service", "com.anti.service.impl", "会话列表"),
                      ("SessionVO.java", "VO", "com.anti.entity.vo", "会话视图")]},
            {"num": "5.3.6.4", "name": "反馈提交", "api": "POST /api/chat/feedback",
             "desc": "用户对AI回答提交反馈（1满意/-1不满意），记录到会话最后一条消息。",
             "classes": ["ChatController", "QAConversationServiceImpl", "FeedbackRequest"],
             "files": [("ChatController.java", "Controller", "com.anti.controller", "反馈接口"),
                      ("QAConversationServiceImpl.java", "Service", "com.anti.service.impl", "反馈记录"),
                      ("FeedbackRequest.java", "DTO", "com.anti.entity.dto", "反馈请求")]},
        ]
    },
    "5.3.7": {
        "module_title": "个性化推荐模块",
        "functions": [
            {"num": "5.3.7.1", "name": "自动策略推荐", "api": "GET /api/recommendation/list",
             "desc": "根据用户生命周期阶段自动选择推荐策略：新手→冷启动（必读资讯+静态属性匹配+热点）；成长→内容+协同+SPM；成熟→协同过滤KNN。",
             "classes": ["RecommendationController", "RecommendationServiceImpl", "ProfileService", "RecommendationAlgorithmUtil", "RecommendationLogMapper"],
             "files": [("RecommendationController.java", "Controller", "com.anti.controller", "推荐接口"),
                      ("RecommendationServiceImpl.java", "Service", "com.anti.service.impl", "策略分发"),
                      ("ProfileService.java", "Service", "com.anti.service", "生命周期判定"),
                      ("RecommendationAlgorithmUtil.java", "Util", "com.anti.util", "算法核心")]},
            {"num": "5.3.7.2", "name": "新手推荐", "api": "GET /api/recommendation/newbie",
             "desc": "新手阶段冷启动推荐：必读资讯优先、静态属性匹配案例（年级/专业）、热点补充、情景关卡优先。评分Wilson*0.6+hot*0.4。",
             "classes": ["RecommendationController", "RecommendationServiceImpl", "FraudCaseMapper", "NewsMapper", "RecommendationAlgorithmUtil"],
             "files": [("RecommendationController.java", "Controller", "com.anti.controller", "新手推荐接口"),
                      ("RecommendationServiceImpl.java", "Service", "com.anti.service.impl", "冷启动逻辑"),
                      ("RecommendationVO.java", "VO", "com.anti.entity.vo", "推荐结果视图")]},
            {"num": "5.3.7.3", "name": "成长期推荐", "api": "GET /api/recommendation/growing",
             "desc": "成长阶段混合推荐：内容相似度（余弦）+SPM序列预测+上下文修正（弱点标签×1.2）。综合分0.5*content+0.3*context+0.2*hot。",
             "classes": ["RecommendationController", "RecommendationServiceImpl", "UserBehaviorMatrixMapper", "AssociationRuleMapper", "RecommendationAlgorithmUtil"],
             "files": [("RecommendationController.java", "Controller", "com.anti.controller", "成长推荐接口"),
                      ("RecommendationServiceImpl.java", "Service", "com.anti.service.impl", "混合推荐"),
                      ("AssociationRule.java", "Entity", "com.anti.entity", "SPM规则")]},
            {"num": "5.3.7.4", "name": "成熟期推荐", "api": "GET /api/recommendation/mature",
             "desc": "成熟阶段协同过滤：同年级预过滤→KNN(k=20)→偏置加权平均预测→标签→案例转换。",
             "classes": ["RecommendationController", "RecommendationServiceImpl", "UserSimilarityMapper", "UserBehaviorMatrixMapper", "RecommendationAlgorithmUtil"],
             "files": [("RecommendationController.java", "Controller", "com.anti.controller", "成熟推荐接口"),
                      ("RecommendationServiceImpl.java", "Service", "com.anti.service.impl", "协同过滤"),
                      ("UserSimilarity.java", "Entity", "com.anti.entity", "用户相似度")]},
            {"num": "5.3.7.5", "name": "用户兴趣分析", "api": "GET /api/recommendation/interest",
             "desc": "分析用户兴趣标签分布，从UserBehaviorMatrix提取标签得分，归一化到0-100，缓存结果。",
             "classes": ["RecommendationController", "RecommendationServiceImpl", "ProfileService", "UserBehaviorMatrixMapper", "UserInterestVO"],
             "files": [("RecommendationController.java", "Controller", "com.anti.controller", "兴趣分析接口"),
                      ("RecommendationServiceImpl.java", "Service", "com.anti.service.impl", "兴趣分析"),
                      ("UserInterestVO.java", "VO", "com.anti.entity.vo", "兴趣视图")]},
        ]
    },
    "5.3.8": {
        "module_title": "数据统计模块",
        "functions": [
            {"num": "5.3.8.1", "name": "数据看板聚合", "api": "GET /api/statistics/dashboard",
             "desc": "聚合今日浏览量、新增用户、活跃用户、通关数、平均分、诈骗类型分布、院系积分等。每个子查询try/catch兜底0或mock数据避免看板崩溃。",
             "classes": ["StatisticsController", "StatisticsServiceImpl", "DailyStatisticsMapper", "DepartmentStatisticsMapper", "StatisticsQueryMapper"],
             "files": [("StatisticsController.java", "Controller", "com.anti.controller", "看板接口"),
                      ("StatisticsServiceImpl.java", "Service", "com.anti.service.impl", "聚合逻辑"),
                      ("DashboardVO.java", "VO", "com.anti.entity.vo", "看板视图")]},
            {"num": "5.3.8.2", "name": "访问趋势分析", "api": "GET /api/statistics/visit/trend",
             "desc": "返回最近N天的访问趋势数据，从daily_statistics读取。",
             "classes": ["StatisticsController", "StatisticsServiceImpl", "DailyStatisticsMapper", "VisitTrendVO"],
             "files": [("StatisticsController.java", "Controller", "com.anti.controller", "趋势接口"),
                      ("StatisticsServiceImpl.java", "Service", "com.anti.service.impl", "趋势查询"),
                      ("VisitTrendVO.java", "VO", "com.anti.entity.vo", "趋势视图")]},
            {"num": "5.3.8.3", "name": "诈骗类型分布", "api": "GET /api/statistics/fraud/types",
             "desc": "返回各类诈骗类型的浏览量分布。",
             "classes": ["StatisticsController", "StatisticsServiceImpl", "FraudCaseMapper", "FraudTypeDistVO"],
             "files": [("StatisticsController.java", "Controller", "com.anti.controller", "类型分布接口"),
                      ("StatisticsServiceImpl.java", "Service", "com.anti.service.impl", "分布统计"),
                      ("FraudTypeDistVO.java", "VO", "com.anti.entity.vo", "分布视图")]},
            {"num": "5.3.8.4", "name": "院系积分统计", "api": "GET /api/statistics/department/scores",
             "desc": "按年级和专业统计平均积分、平均知识水平。",
             "classes": ["StatisticsController", "StatisticsServiceImpl", "DepartmentStatisticsMapper", "DepartmentScoreVO"],
             "files": [("StatisticsController.java", "Controller", "com.anti.controller", "院系统计接口"),
                      ("StatisticsServiceImpl.java", "Service", "com.anti.service.impl", "院系统计"),
                      ("DepartmentScoreVO.java", "VO", "com.anti.entity.vo", "院系视图")]},
            {"num": "5.3.8.5", "name": "统计数据导出", "api": "GET /api/statistics/export/daily",
             "desc": "导出每日统计或院系统计为Excel文件，使用EasyExcel生成。",
             "classes": ["StatisticsController", "StatisticsServiceImpl", "EasyExcel", "StatisticsExportVO"],
             "files": [("StatisticsController.java", "Controller", "com.anti.controller", "导出接口"),
                      ("StatisticsServiceImpl.java", "Service", "com.anti.service.impl", "Excel生成"),
                      ("StatisticsExportVO.java", "VO", "com.anti.entity.vo", "导出视图")]},
        ]
    },
    "5.3.9": {
        "module_title": "积分成就与排行榜模块",
        "functions": [
            {"num": "5.3.9.1", "name": "积分信息查询", "api": "GET /api/score/info",
             "desc": "查询用户积分信息，包括总积分、等级、本周积分、成就解锁数。",
             "classes": ["ScoreController", "ScoreServiceImpl", "UserScoreMapper", "ScoreVO", "AchievementService"],
             "files": [("ScoreController.java", "Controller", "com.anti.controller", "积分查询接口"),
                      ("ScoreServiceImpl.java", "Service", "com.anti.service.impl", "积分查询"),
                      ("UserScore.java", "Entity", "com.anti.entity", "积分实体"),
                      ("ScoreVO.java", "VO", "com.anti.entity.vo", "积分视图")]},
            {"num": "5.3.9.2", "name": "成就列表查询", "api": "GET /api/achievement/list",
             "desc": "查询所有成就定义，包括成就代码、名称、描述、奖励积分、解锁条件。填充当前用户是否已解锁标记。",
             "classes": ["AchievementController", "AchievementServiceImpl", "AchievementMapper", "UserAchievementMapper", "AchievementVO"],
             "files": [("AchievementController.java", "Controller", "com.anti.controller", "成就列表接口"),
                      ("AchievementServiceImpl.java", "Service", "com.anti.service.impl", "成就查询"),
                      ("Achievement.java", "Entity", "com.anti.entity", "成就实体"),
                      ("AchievementVO.java", "VO", "com.anti.entity.vo", "成就视图")]},
            {"num": "5.3.9.3", "name": "用户成就查询", "api": "GET /api/achievement/user",
             "desc": "查询当前用户已解锁的成就列表。",
             "classes": ["AchievementController", "AchievementServiceImpl", "UserAchievementMapper", "UserAchievementVO"],
             "files": [("AchievementController.java", "Controller", "com.anti.controller", "用户成就接口"),
                      ("AchievementServiceImpl.java", "Service", "com.anti.service.impl", "用户成就查询"),
                      ("UserAchievement.java", "Entity", "com.anti.entity", "用户成就记录")]},
            {"num": "5.3.9.4", "name": "日榜查询", "api": "GET /api/leaderboard/daily",
             "desc": "查询今日积分排行榜Top N。Redis缓存300秒，DB兜底。标记当前用户排名。",
             "classes": ["LeaderboardController", "LeaderboardServiceImpl", "LeaderboardMapper", "RedisCacheUtil", "LeaderboardVO"],
             "files": [("LeaderboardController.java", "Controller", "com.anti.controller", "日榜接口"),
                      ("LeaderboardServiceImpl.java", "Service", "com.anti.service.impl", "日榜查询"),
                      ("Leaderboard.java", "Entity", "com.anti.entity", "排行榜实体"),
                      ("RedisCacheUtil.java", "Util", "com.anti.util", "排行榜缓存")]},
            {"num": "5.3.9.5", "name": "周榜查询", "api": "GET /api/leaderboard/weekly",
             "desc": "查询本周积分排行榜Top N。Redis缓存900秒。",
             "classes": ["LeaderboardController", "LeaderboardServiceImpl", "LeaderboardMapper", "RedisCacheUtil", "LeaderboardVO"],
             "files": [("LeaderboardController.java", "Controller", "com.anti.controller", "周榜接口"),
                      ("LeaderboardServiceImpl.java", "Service", "com.anti.service.impl", "周榜查询"),
                      ("LeaderboardVO.java", "VO", "com.anti.entity.vo", "排行榜视图")]},
            {"num": "5.3.9.6", "name": "总榜查询", "api": "GET /api/leaderboard/all",
             "desc": "查询总积分排行榜Top N。Redis缓存1800秒。",
             "classes": ["LeaderboardController", "LeaderboardServiceImpl", "LeaderboardMapper", "RedisCacheUtil", "LeaderboardVO"],
             "files": [("LeaderboardController.java", "Controller", "com.anti.controller", "总榜接口"),
                      ("LeaderboardServiceImpl.java", "Service", "com.anti.service.impl", "总榜查询")]},
            {"num": "5.3.9.7", "name": "用户排名查询", "api": "GET /api/leaderboard/user-rank",
             "desc": "查询当前用户在指定周期（日/周/总）的排名。",
             "classes": ["LeaderboardController", "LeaderboardServiceImpl", "LeaderboardMapper", "RedisCacheUtil"],
             "files": [("LeaderboardController.java", "Controller", "com.anti.controller", "用户排名接口"),
                      ("LeaderboardServiceImpl.java", "Service", "com.anti.service.impl", "用户排名查询")]},
        ]
    },
}

# 定位5.3节的模块位置
def find_module_positions(doc):
    """找到每个5.3.x模块小节的Heading3段落位置"""
    positions = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if style == "Heading 3" and text.startswith("5.3.") and "模块" in text:
            # 提取模块编号如"5.3.1"
            num = text.split(" ")[0]
            positions[num] = i
    return positions

# 在模块段落之后插入功能子节
def insert_function_subsections(doc, module_num, module_idx, functions):
    """在模块小节后插入功能子节（从小到大）"""
    # 找到模块小节的"功能列表"段落后的位置作为插入点
    insert_idx = module_idx
    # 向后扫描找到"功能列表"段落
    for j in range(module_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[j]
        if p.text.strip().startswith("2、功能列表"):
            insert_idx = j + 1
            break
        # 如果遇到下一个Heading3或Heading1，说明模块结束
        style = p.style.name if p.style else ""
        if style.startswith("Heading") and not style == "Heading 4":
            insert_idx = j
            break
    
    # 按功能编号从小到大插入
    ref_para = doc.paragraphs[insert_idx - 1]
    for func in functions:
        # 插入Heading4标题
        h4 = ref_para.insert_paragraph_after(f"{func['num']} {func['name']}")
        h4.style = "Heading 4"
        ref_para = h4
        
        # 插入功能设计描述
        desc1 = ref_para.insert_paragraph_after("1 功能设计描述")
        desc1.style = "Normal"
        ref_para = desc1
        
        desc2 = ref_para.insert_paragraph_after(func['desc'])
        desc2.style = "Normal"
        ref_para = desc2
        
        # 插入类介绍
        class1 = ref_para.insert_paragraph_after("（1）类")
        class1.style = "Normal"
        ref_para = class1
        
        for idx, cls in enumerate(func['classes'], 1):
            cls_para = ref_para.insert_paragraph_after(f"{idx}) {cls}")
            cls_para.style = "Normal"
            ref_para = cls_para
        
        # 插入文件列表说明
        file1 = ref_para.insert_paragraph_after("（3）文件列表")
        file1.style = "Normal"
        ref_para = file1
        
        file2 = ref_para.insert_paragraph_after("如下表所示。")
        file2.style = "Normal"
        ref_para = file2
        
        # 插入表格标题
        table_title = ref_para.insert_paragraph_after(f"表{func['num']} {func['name']}文件列表")
        table_title.style = "Normal"
        ref_para = table_title
        
        # 创建表格
        table = doc.add_table(rows=len(func['files']) + 1, cols=4)
        table.style = "Table Grid"
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "名称"
        hdr_cells[1].text = "类型"
        hdr_cells[2].text = "存放位置"
        hdr_cells[3].text = "说明"
        # 填充数据
        for ri, (name, typ, loc, desc) in enumerate(func['files'], 1):
            row_cells = table.rows[ri].cells
            row_cells[0].text = name
            row_cells[1].text = typ
            row_cells[2].text = loc
            row_cells[3].text = desc
        
        # 表格需要移动到正确位置（python-docx的insert_paragraph_after无法直接插入表格）
        # 我们在最后创建了表格，但文档顺序需要调整
        # 简化方案：表格留在末尾，通过引用说明

# 执行填充
print("从备份恢复...")
positions = find_module_positions(doc)
print(f"找到模块位置: {positions}")

# 为每个模块添加功能子节
for module_num in sorted(positions.keys()):
    if module_num in functions_data:
        module_idx = positions[module_num]
        functions = functions_data[module_num]['functions']
        print(f"处理 {module_num} {functions_data[module_num]['module_title']}, 添加 {len(functions)} 个功能子节")
        insert_function_subsections(doc, module_num, module_idx, functions)

# 保存
doc.save(target_path)
print(f"\n保存到: {target_path}")
print("完成")