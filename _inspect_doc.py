#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查Word文档中的所有Heading 4标题
"""

from pathlib import Path
from docx import Document

BASE_DIR = Path(r"D:\Project\anti_fraud_platform")
TARGET_FILE = BASE_DIR / "docx" / "《项目阶段二》评审" / "专业综合课程设计-系统设计说明书-江乐霖-23201317.docx"

def main():
    doc = Document(TARGET_FILE)
    
    print("所有Heading 4标题：")
    print("=" * 60)
    
    count = 0
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 4':
            count += 1
            print(f"{count}. [{i}] {para.text}")
    
    print("=" * 60)
    print(f"总计：{count} 个Heading 4标题")
    
    # 检查5.3节模块标题
    print("\n5.3节模块标题（Heading 3）：")
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 3' and '5.3' in para.text:
            print(f"  [{i}] {para.text}")


if __name__ == '__main__':
    main()