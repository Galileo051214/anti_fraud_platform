# -*- coding: utf-8 -*-
"""
生成第六章-用户界面设计概述独立文档
内容与第六章_界面原型文字描述.md一致，按A/B/C格式输出
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

REF_PATH = r"D:\Project\anti_fraud_platform\docx\《项目阶段二》评审\[2024]3-系统设计.docx"
OUTPUT_DIR = r"D:\Project\anti_fraud_platform\docx\《项目阶段二》评审"
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "第六章-用户界面设计概述.docx")
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document(REF_PATH)
body = doc.element.body
for element in list(body):
    if element.tag.endswith('sectPr'):
        continue
    body.remove(element)


def set_font(run, size_pt, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 16, True)
    p.style = 'Heading 1'
    p.paragraph_format.space_before = Pt(11.7)
    p.paragraph_format.space_after = Pt(11.7)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 14, True)
    p.style = 'Heading 2'
    p.paragraph_format.space_before = Pt(11.7)
    p.paragraph_format.space_after = Pt(11.7)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 14, True)
    p.style = 'Heading 3'
    p.paragraph_format.space_before = Pt(1.4)
    p.paragraph_format.space_after = Pt(1.4)
    return p

def add_h4(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 12, True)
    p.style = 'Heading 4'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_normal(text, bold=False, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 10.5, bold)
    p.style = 'Normal'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Pt(2)
    return p

def add_bullet(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 10.5, False)
    p.style = 'Normal'
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(0)
    return p

def add_table(caption_text, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = '网格表 4 - 着色 31'
    except:
        table.style = 'Table Grid'
    # header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, 10.5, True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # body
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, 10.5, False)
    # caption
    cp = doc.add_paragraph()
    run = cp.add_run(caption_text)
    set_font(run, 10.5, True)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

def add_section_label(text):
    """添加A/B/C分段标签，如"A. 界面整体功能描述" """
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 10.5, True)
    p.style = 'Normal'
    p.paragraph_format.first_line_indent = Pt(2)
    return p


# ========================== 开始构建文档 ==========================

add_h1("6 用户界面设计概述")
add_normal("用户界面设计是系统与用户之间交互的桥梁，良好的界面设计能够提升用户体验、降低学习成本。本章从整体设计原则、布局规范、配色方案等方面阐述反诈骗学习平台的界面设计思路，并给出各主要功能模块的界面原型文字描述。每个界面按照 A.界面整体功能描述、B.数据描述、C.按键功能三个维度详细说明。")

add_h2("6.1 用户界面设计")
add_normal("本节从设计原则、布局规范、配色方案、响应式设计以及交互反馈五个方面，描述反诈骗学习平台的整体界面设计风格与规范。")

# 6.1.1 登录与注册模块
add_h3("6.1.1 登录与注册模块")

add_h4("6.1.1.1 登录界面（VC101）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：已注册的学生、管理员")
add_bullet("主要功能：输入用户名和密码进行身份认证，校验通过后获取 JWT 令牌并跳转至对应首页")
add_bullet("完成后转向界面：学生用户 → /home（学生端首页）；管理员用户 → /admin/dashboard（管理端数据看板）")
add_bullet("无账号处理：提供「立即注册」链接跳转注册界面；提供「忘记密码」链接跳转找回密码界面（预留）")

add_section_label("B. 数据描述")
add_table("表6.1.1.1 登录界面数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["用户名", "VARCHAR(50)", "3~20字符，支持字母、数字、下划线", "用户在注册时填写，唯一索引", "登录界面仅输入，不可编辑已有数据"],
    ["密码", "VARCHAR(128)", "至少6字符，前端明文传输，后端BCrypt加密存储", "用户在注册时设置", "登录界面仅输入，不可编辑已有数据"],
    ["记住密码状态", "BOOLEAN", "true/false", "用户勾选「记住密码」复选框决定", "登录前可选"],
])
add_bullet("校验规则：用户名必填，3~20字符，正则/^[a-zA-Z0-9_]{3,20}$/；密码必填，长度≥6")
add_bullet("校验失败时输入框下方显示红色错误提示文字")

add_section_label("C. 按键功能")
add_table("表6.1.1.2 登录界面按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["登录", "提交认证请求", "前端校验用户名和密码非空及格式→调用POST /api/user/login，Body为{username, password}→成功：后端返回{token, role}，前端将token存入localStorage，根据role跳转/home或/admin/dashboard；失败：显示「用户名或密码错误」提示"],
    ["注册链接", "跳转注册页面", "无校验，直接路由跳转至/register"],
    ["忘记密码", "跳转找回密码页面", "预留功能，跳转至/forgot-password"],
])

add_h4("6.1.1.2 注册界面（VC102）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：游客（未注册用户）")
add_bullet("主要功能：填写个人信息完成账号注册，注册成功后系统自动创建用户、初始化积分（0）和用户画像（空），转入登录界面")
add_bullet("完成后转向界面：登录界面（VC101）")

add_section_label("B. 数据描述")
add_table("表6.1.1.3 注册界面数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["用户名", "VARCHAR(50)", "3~20字符，字母/数字/下划线，唯一", "用户输入，后端校验唯一性", "注册后不可修改"],
    ["学号", "VARCHAR(30)", "唯一，可包含字母和数字", "用户输入，后端校验唯一性", "注册后不可修改"],
    ["密码", "VARCHAR(128)", "≥6字符，BCrypt加密存储", "用户输入，前端与确认密码比对一致后传输", "注册后可修改（通过个人中心）"],
    ["确认密码", "—", "与密码一致", "用户输入，前端比对", "—"],
    ["昵称", "VARCHAR(50)", "≤50字符，可空", "用户输入，为空则默认取用户名", "个人中心可修改"],
    ["手机号", "VARCHAR(20)", "11位数字，合法手机号格式", "用户输入", "个人中心可修改"],
    ["邮箱", "VARCHAR(100)", "合法邮箱格式（xxx@xxx.xxx）", "用户输入", "个人中心可修改"],
    ["年级", "VARCHAR(20)", "枚举值：大一、大二、大三、大四、研一、研二、研三", "用户从下拉框选择", "个人中心可修改"],
    ["专业", "VARCHAR(100)", "≤100字符，可空", "用户输入", "个人中心可修改"],
    ["角色", "ENUM('student','admin')", "默认'student'", "系统自动生成", "不可由用户修改"],
    ["状态", "TINYINT", "0=禁用，1=正常（默认）", "系统自动生成", "管理员后台管理"],
])
add_bullet("校验规则：用户名必填（3~20字符，唯一正则校验）；学号必填（唯一）；密码必填（≥6字符）；确认密码必填（必须与密码完全一致，否则提示「两次密码输入不一致」）；手机号选填（若填写则校验11位数字格式）；邮箱选填（若填写则校验邮箱正则）")

add_section_label("C. 按键功能")
add_table("表6.1.1.4 注册界面按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["注册", "提交注册请求", "前端逐字段校验格式和必填项→确认密码与密码一致性校验→调用POST /api/user/register，Body为{username, password, studentNo, nickname, phone, email, grade, major}→成功：后端返回成功消息，前端弹窗提示「注册成功」并跳转至登录页→失败（如用户名或学号已存在）：显示具体错误信息"],
    ["返回登录", "跳转登录页面", "无校验，直接路由跳转至/login"],
])
add_bullet("后端注册逻辑：密码使用BCrypt加密后存储；自动初始化user_score表（积分=0）；自动初始化user_profile表（画像为空JSON{}）；写入user表，create_time取数据库当前时间")

# 6.1.2 首页模块
add_h3("6.1.2 首页模块")

add_h4("6.1.2.1 首页界面（VC201）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：所有用户（包括未登录游客和已登录学生/管理员）")
add_bullet("主要功能：展示平台核心信息，包括Hero宣传区（平台口号）、平台特色功能区（游戏化学习/智能客服/智能推荐/社区互动）、热门案例推荐区")
add_bullet("完成后转向界面：首页作为主入口页面，各操作按钮可跳转至对应功能模块，无「完成」后的固定转向")

add_section_label("B. 数据描述")
add_table("表6.1.2.1 首页数据描述", ["数据项", "类型", "格式/说明", "生成方式", "是否允许编辑"], [
    ["平台口号", "VARCHAR(100)", "固定文本", "前端硬编码", "否"],
    ["热门案例列表", "JSON Array", "[{id, title, type, views, likes}]", "调用GET /api/case/hot获取，后端按浏览量降序返回TOP6", "否（只读展示）"],
])
add_bullet("静态内容：Hero标题、副标题、特色功能卡片文案→前端硬编码")

add_section_label("C. 按键功能")
add_table("表6.1.2.2 首页按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["「开始学习」", "跳转知识闯关", "路由跳转至/challenge"],
    ["「了解诈骗案例」", "跳转案例库", "路由跳转至/case"],
    ["特色功能卡片", "跳转对应模块", "四张卡片分别跳转：闯关(/challenge)、客服(/chat)、推荐(/recommend)、社区(/forum)"],
    ["热门案例卡片", "跳转案例详情", "点击后路由跳转至/case/:id，:id为该案例的数据库主键"],
])

# 6.1.3 资讯学习模块
add_h3("6.1.3 资讯学习模块")

add_h4("6.1.3.1 资讯列表界面（VC301）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：以卡片网格形式展示资讯列表，支持按分类筛选、按类型（新闻/预警/政策）筛选、关键词搜索，分页加载")
add_bullet("完成后转向界面：点击资讯卡片→跳转资讯详情界面（VC302），分页操作不离开本界面")

add_section_label("B. 数据描述")
add_table("表6.1.3.1 资讯列表数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["分类列表", "JSON Array", "[{id, name}]", "调用GET /api/news/categories获取", "否（只读）"],
    ["资讯ID", "BIGINT", "自增主键", "系统自动生成", "否"],
    ["标题", "VARCHAR(200)", "≤200字符", "管理员后台发布时录入", "管理员可编辑"],
    ["类型", "TINYINT", "0=新闻，1=预警，2=政策", "管理员发布时选择", "管理员可编辑"],
    ["置顶/必读", "BOOLEAN", "true/false", "管理员发布时设置", "管理员可编辑"],
    ["浏览量", "INT", "非负整数", "每次用户进入详情页时+1", "否（系统自动）"],
    ["点赞数", "INT", "非负整数", "用户点击点赞按钮时+1/-1", "否（系统自动）"],
    ["发布时间", "DATETIME", "yyyy-MM-dd HH:mm:ss", "管理员发布时系统自动生成", "否"],
    ["分页参数", "—", "pageNum(默认1), pageSize(默认10)", "用户操作切换", "动态查询"],
])
add_bullet("查询逻辑：分类筛选→URL参数category，精确匹配；类型筛选→URL参数type，精确匹配；关键词搜索→URL参数keyword，模糊匹配标题和内容，支持中文分词")

add_section_label("C. 按键功能")
add_table("表6.1.3.2 资讯列表按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["分类下拉框", "筛选分类", "选择后重新加载列表，category参数变更"],
    ["类型标签", "筛选类型", "点击「新闻/预警/政策」，type参数变更，高亮当前选中"],
    ["搜索按钮", "关键词搜索", "输入关键词后按回车或点击搜索图标，keyword参数变更"],
    ["资讯卡片", "跳转详情", "点击触发路由/news/:id"],
    ["点赞按钮", "点赞/取消点赞", "调用POST /api/news/{id}/like，已点赞则取消（toggle），乐观更新"],
    ["分页器", "切换页码", "点击页码或「上一页/下一页」，pageNum参数变更"],
])

add_h4("6.1.3.2 资讯详情界面（VC302）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：展示单篇资讯的完整内容，含标题、作者、发布时间、浏览量、富文本正文、点赞功能和相关资讯推荐")
add_bullet("完成后转向界面：无固定转向，可点击「返回」回列表，或点击相关资讯跳转其他详情")

add_section_label("B. 数据描述")
add_table("表6.1.3.3 资讯详情数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["标题", "VARCHAR(200)", "≤200字符", "管理员发布", "否（只读）"],
    ["作者", "VARCHAR(50)", "管理员用户名", "从发布者的token中获取", "否"],
    ["发布时间", "DATETIME", "yyyy-MM-dd HH:mm:ss", "自动", "否"],
    ["浏览量", "INT", "非负", "用户进入详情页时+1", "否（自动）"],
    ["点赞数", "INT", "非负", "用户点击点赞时更新", "否（自动）"],
    ["内容", "TEXT/LONGTEXT", "富文本HTML", "管理员发布", "否（只读）"],
    ["封面图", "VARCHAR(255)", "图片URL", "管理员上传", "否"],
    ["相关资讯", "JSON Array", "[{id, title}]", "后端按同类/同标签推荐3-5条", "否"],
])
add_bullet("数据上报：页面加载时调用POST /api/news/{id}/browse记录浏览（增量+1）；页面关闭或路由离开时上报停留时长（秒），用于用户行为分析")

add_section_label("C. 按键功能")
add_table("表6.1.3.4 资讯详情按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["← 返回列表", "返回到资讯列表", "router.back()或路由跳转至/news"],
    ["点赞按钮", "点赞/取消", "调用POST /api/news/{id}/like，toggle模式"],
    ["相关资讯卡片", "跳转相关资讯", "路由/news/:relatedId"],
])

# 6.1.4 案例展示模块
add_h3("6.1.4 案例展示模块")

add_h4("6.1.4.1 案例列表界面（VC401）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：以卡片网格展示诈骗案例库，支持关键词搜索和标签筛选。卡片展示案例标题、类型标签、风险等级徽章、难度等级、浏览量和点赞数")
add_bullet("完成后转向界面：点击案例卡片→跳转案例详情界面（VC402）")

add_section_label("B. 数据描述")
add_table("表6.1.4.1 案例列表数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["案例ID", "BIGINT", "自增主键", "系统自动", "否"],
    ["标题", "VARCHAR(200)", "≤200字符", "管理员录入", "管理员可编辑"],
    ["类型", "TINYINT", "枚举：0=刷单，1=冒充，2=校园贷，3=投资理财，4=裸聊敲诈，5=其他", "管理员选择", "管理员可编辑"],
    ["标签", "JSON", "[\"标签1\",\"标签2\"]", "管理员录入", "管理员可编辑"],
    ["难度", "TINYINT", "1~5整数", "管理员设置", "管理员可编辑"],
    ["风险等级", "ENUM", "极高/高/中/低/极低", "管理员设置", "管理员可编辑"],
    ["浏览量", "INT", "非负", "用户访问详情时+1", "否（自动）"],
    ["点赞数", "INT", "非负", "用户点赞时+1/-1", "否（自动）"],
    ["发布时间", "DATETIME", "yyyy-MM-dd HH:mm:ss", "自动", "否"],
    ["精选标识", "BOOLEAN", "true/false", "管理员标记", "管理员可编辑"],
])
add_bullet("查询逻辑：关键词搜索→模糊匹配标题和内容；标签筛选→点击标签后作为筛选条件，支持组合筛选")

add_section_label("C. 按键功能")
add_table("表6.1.4.2 案例列表按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["搜索框", "关键词搜索", "输入后按回车，调用GET /api/case/list?keyword=xxx，支持模糊搜索"],
    ["标签云标签", "按标签筛选", "点击标签切换筛选，标签高亮显示，再次点击取消筛选"],
    ["案例卡片", "跳转详情", "路由跳转至/case/:id"],
    ["点赞按钮", "点赞/取消", "调用POST /api/case/{id}/like，toggle模式，需登录"],
])

add_h4("6.1.4.2 案例详情界面（VC402）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：展示诈骗案例完整详情，包含剧情内容（富文本）、FSM诈骗剧本节点可视化、风险提示仪表盘、适用对象和相关案例推荐")
add_bullet("完成后转向界面：无固定转向，可返回列表或点击相关案例跳转")

add_section_label("B. 数据描述")
add_table("表6.1.4.3 案例详情数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["内容", "TEXT/LONGTEXT", "富文本HTML", "管理员录入", "管理员可编辑"],
    ["诈骗剧本", "JSON", "{nodes:[{id,type,label,content}], edges:[{from,to,label,isSafeChoice}]}", "管理员录入FSM状态机定义", "管理员可编辑"],
    ["适用对象", "VARCHAR(200)", "如「大一新生、所有学生」", "管理员设置", "管理员可编辑"],
    ["风险指数", "INT", "0~10整数", "管理员设置或由风险等级映射", "管理员可编辑"],
    ["相关案例列表", "JSON Array", "[{id, title}]", "后端按类型/标签推荐4-6条", "否"],
])
add_bullet("数据展示格式：风险仪表盘→圆形进度条，颜色随风险指数变化（0~3绿色，4~6橙色，7~10红色）；FSM剧本→以流程图形式展示节点（start→dialog→decision→result→end），安全选择边为绿色，危险选择边为红色")

add_section_label("C. 按键功能")
add_table("表6.1.4.4 案例详情按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["← 返回列表", "返回案例列表", "router.back()或跳转/case"],
    ["点赞按钮", "点赞/取消", "调用POST /api/case/{id}/like，toggle模式"],
    ["相关案例卡片", "跳转相关案例", "路由/case/:relatedId"],
    ["进入情景模拟", "跳转情景模拟（如有）", "路由跳转至/challenge/scenario/:id"],
])

# 6.1.5 知识闯关模块
add_h3("6.1.5 知识闯关模块")

add_h4("6.1.5.1 关卡列表界面（VC501）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：展示所有知识闯关关卡，包含总体进度统计（已通关数/总关卡数/完成率）、类型筛选（全部/答题挑战/情景模拟）和关卡卡片网格。卡片展示关卡序号、类型徽章、状态标签（已通关/待挑战/已锁定）、标题、难度、积分奖励")
add_bullet("完成后转向界面：点击关卡卡片→答题挑战界面（VC502）或情景模拟界面（VC503）")

add_section_label("B. 数据描述")
add_table("表6.1.5.1 关卡列表数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["关卡ID", "BIGINT", "自增主键", "系统自动", "否"],
    ["标题", "VARCHAR(200)", "≤200字符", "管理员录入", "管理员可编辑"],
    ["描述", "TEXT", "关卡背景说明", "管理员录入", "管理员可编辑"],
    ["类型", "ENUM", "'quiz'(答题挑战)/'scenario'(情景模拟)", "管理员设置", "管理员可编辑"],
    ["难度", "TINYINT", "1~5整数", "管理员设置", "管理员可编辑"],
    ["积分奖励", "INT", "正整数", "管理员设置", "管理员可编辑"],
    ["通关状态", "—", "已通关/待挑战/已锁定", "根据前置关卡完成情况自动计算", "可根据重置需求手动重置"],
])
add_bullet("状态计算规则：第N关通关状态=已通关（得分≥60%）；第N关未通关且第N-1关已通关或N=1=待挑战；第N-1关未通关且N>1=已锁定")

add_section_label("C. 按键功能")
add_table("表6.1.5.2 关卡列表按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["类型标签页（全部/答题挑战/情景模拟）", "筛选关卡", "切换后仅显示对应类型的关卡，默认「全部」"],
    ["关卡卡片", "进入挑战", "已通关：可重新挑战；待挑战：进入答题/情景模拟；已锁定：提示「需先完成上一关卡」"],
    ["进度统计区", "—", "只读展示"],
])

add_h4("6.1.5.2 答题挑战界面（VC502）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：逐题展示答题界面，支持单选题和多选题。顶部显示答题进度条，中部为题目卡片（题型标签、分值、题干、选项列表），底部为导航按钮。答题完成后展示成绩评级（S/A/B/C/D）、正确统计和通关信息")
add_bullet("完成后转向界面：提交后停留在成绩页面→点击「返回关卡」回VC501")

add_section_label("B. 数据描述")
add_table("表6.1.5.3 答题挑战数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["题目ID", "BIGINT", "自增主键", "系统自动", "否"],
    ["题型", "ENUM", "'single'(单选)/'multiple'(多选)", "管理员设置", "管理员可编辑"],
    ["题干", "TEXT", "题目文字", "管理员录入", "管理员可编辑"],
    ["选项列表", "JSON", "[{label:'A', text:'...'}]", "管理员录入", "管理员可编辑"],
    ["正确答案", "VARCHAR(10)", "如'A'或'ABC'", "管理员设置", "管理员可编辑"],
    ["分值", "INT", "正整数", "管理员设置", "管理员可编辑"],
    ["用户答案", "VARCHAR(10)", "用户选择的选项", "用户答题时记录", "答题过程中可修改"],
])
add_bullet("评级规则：≥90%→S（完美通关），≥80%→A（优秀），≥70%→B（良好），≥60%→C（及格），<60%→D（未通过）")
add_bullet("积分获得：积分=基础奖励×(正确率/100%)，S评级额外获得+5积分")

add_section_label("C. 按键功能")
add_table("表6.1.5.4 答题挑战按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["选项卡片", "选择答案", "单击选中/取消，单选模式自动互斥，多选模式可多选"],
    ["上一题", "切换到上一题", "保留已选答案，到第一题时禁用"],
    ["下一题", "切换到下一题", "需确认已选答案（至少选择一个），到最后一道题时变为「提交」"],
    ["提交", "提交所有答案（最后一道题）", "前端校验所有题目是否已回答→调用POST /api/challenge/submit，Body为{challengeId, answers:[{questionId, selected}]}→后端自动评分并返回结果{score, correctCount, totalCount, grade, passed, earnedPoints}→前端展示成绩动画和评级"],
    ["再试一次", "重置答题", "清空所有答案，重新从第一题开始"],
    ["返回关卡", "回关卡列表", "路由跳转至/challenge"],
])

add_h4("6.1.5.3 情景模拟界面（VC503）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：以对话式交互模拟真实诈骗场景。FSM（有限状态机）驱动剧情流转：用户阅读角色对话（诈骗分子/受害者/旁白），在每个决策点做出选择，系统根据选择判定是否安全并推进剧情。结束后展示评分、评级、正确决策数和获得积分")
add_bullet("完成后转向界面：点击「返回关卡」回VC501")

add_section_label("B. 数据描述")
add_table("表6.1.5.5 情景模拟数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["节点ID", "VARCHAR(50)", "如'node_001'", "管理员在FSM编辑器中定义", "管理员可编辑"],
    ["节点类型", "ENUM", "'start'/'dialog'/'decision'/'result'/'end'", "管理员定义", "管理员可编辑"],
    ["角色", "ENUM", "'narrator'(旁白)/'victim'(受害者)/'scammer'(诈骗分子)", "管理员定义", "管理员可编辑"],
    ["对话内容", "TEXT", "角色台词", "管理员录入", "管理员可编辑"],
    ["可选边列表", "JSON", "[{edgeId, label, condition, isSafeChoice, targetNodeId}]", "管理员定义FSM边", "管理员可编辑"],
    ["风险提示", "TEXT", "如「此选择存在风险！」", "管理员在决策节点配置", "管理员可编辑"],
])
add_bullet("决策历史记录：用户在每步的选择、时间戳、是否正确，全部记录用于最终评分")
add_bullet("评分逻辑：总分=正确决策数/总决策数×100（百分制），评级与答题挑战一致（S/A/B/C/D）")

add_section_label("C. 按键功能")
add_table("表6.1.5.6 情景模拟按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["选项按钮（决策节点）", "做出决策", "调用POST /api/challenge/decision，Body{scenarioId, nodeId, selectedEdgeId}→后端返回下一个节点数据（对话内容或结果）→前端根据返回更新对话区→若isSafeChoice=false，显示风险提示框"],
    ["重新挑战", "重置情景", "清空决策历史，从start节点重新开始"],
    ["返回关卡", "回关卡列表", "路由跳转至/challenge"],
])

# 6.1.6 社区论坛模块
add_h3("6.1.6 社区论坛模块")

add_h4("6.1.6.1 帖子列表界面（VC601）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：以卡片流形式展示社区帖子，支持按类型标签页筛选（经验分享/求助问答/讨论交流）、排序（时间/点赞/评论）、关键词搜索和发布新帖子")
add_bullet("完成后转向界面：点击帖子卡片→帖子详情界面（VC602）")

add_section_label("B. 数据描述")
add_table("表6.1.6.1 帖子列表数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["帖子ID", "BIGINT", "自增主键", "系统自动", "否"],
    ["标题", "VARCHAR(200)", "≤50字符", "用户发布时录入", "用户可编辑（限本人，1小时内）"],
    ["内容", "TEXT", "帖子正文", "用户发布时录入", "用户可编辑（限本人，1小时内）"],
    ["类型", "ENUM", "'experience'/'question'/'discussion'", "用户发布时选择", "否"],
    ["作者ID", "BIGINT", "外键→user(id)", "从登录token获取", "否"],
    ["浏览量", "INT", "非负", "用户进入详情时+1", "否（自动）"],
    ["评论数", "INT", "非负", "有人发表评论时增加", "否（自动）"],
    ["点赞数", "INT", "非负", "用户点赞时+/-1", "否（自动）"],
    ["发布时间", "DATETIME", "yyyy-MM-dd HH:mm:ss，前端显示为「X小时前」", "自动", "否"],
])
add_bullet("排序逻辑：按时间→ORDER BY create_time DESC；按点赞→ORDER BY like_count DESC；按评论→ORDER BY comment_count DESC")
add_bullet("分页：默认每页10条，支持滚动加载（infinite scroll）")

add_section_label("C. 按键功能")
add_table("表6.1.6.2 帖子列表按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["类型标签页", "筛选帖子类型", "切换后仅显示对应类型的帖子"],
    ["排序下拉框", "切换排序", "按时间/点赞/评论排序，实时刷新"],
    ["搜索框", "搜索帖子", "模糊匹配标题和内容"],
    ["发布帖子", "打开发布弹窗", "弹出一个模态对话框，包含：类型选择（下拉）、标题输入（≤50字，必填）、内容输入（必填）"],
    ["帖子卡片", "跳转详情", "路由跳转至/forum/:id"],
    ["点赞按钮", "点赞/取消", "调用POST /api/forum/{id}/like，toggle模式"],
])
add_bullet("发布帖子弹窗提交校验逻辑：类型必选；标题必填（≤50字符）；内容必填。校验通过→调用POST /api/forum/publish→成功重新加载列表→失败显示错误提示")
add_bullet("删除帖子规则：用户只能删除自己的帖子，删除采用逻辑删除（仅打标记，数据库记录保留），管理员可物理删除违规帖子")

add_h4("6.1.6.2 帖子详情界面（VC602）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：展示帖子完整内容、点赞统计和评论列表。支持发表评论、回复评论、点赞评论、删除自己的评论。评论采用嵌套层级展示（一级评论+子评论）")
add_bullet("完成后转向界面：无固定转向，可返回列表")

add_section_label("B. 数据描述")
add_table("表6.1.6.3 帖子详情数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["帖子内容", "TEXT", "富文本/纯文本", "用户发布时录入", "否（只读）"],
    ["评论ID", "BIGINT", "自增主键", "系统自动", "否"],
    ["评论内容", "TEXT", "≤1000字符", "用户输入", "用户可编辑（限本人，30分钟内）"],
    ["父评论ID", "BIGINT", "可空，为空表示一级评论", "系统自动关联", "否"],
    ["是否作者标识", "BOOLEAN", "评论者=帖子作者时为true", "系统自动判断", "否"],
])
add_bullet("评论嵌套：一级评论直接展示在帖子下方；子评论默认折叠，点击「展开X条回复」后展示；最多支持两级嵌套")

add_section_label("C. 按键功能")
add_table("表6.1.6.4 帖子详情按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["点赞按钮", "点赞帖子", "调用POST /api/forum/{id}/like"],
    ["评论输入框+发表按钮", "发表一级评论", "内容非空校验→调用POST /api/forum/comment，Body{postId, content}→成功后评论列表追加新评论"],
    ["回复按钮（评论下方）", "展开回复输入框", "点击后显示内联输入框，可@回复对象"],
    ["→回复提交", "提交子评论", "调用POST /api/forum/comment，Body{postId, parentId, content}"],
    ["删除按钮（评论右上角）", "删除评论", "仅显示当前登录用户的评论→点击弹出确认对话框「确定删除此评论？」→确定后调用DELETE /api/forum/comment/{commentId}，逻辑删除（打标记）"],
    ["⬆展开/收起子评论", "展开或收起子评论列表", "切换显示状态"],
])

# 6.1.7 个人中心模块
add_h3("6.1.7 个人中心模块")

add_h4("6.1.7.1 个人中心界面（VC701）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：已登录学生")
add_bullet("主要功能：用户管理个人信息、修改密码和查看学习记录的统一入口。左侧边栏展示用户信息卡、统计数据、导航菜单和快捷入口；右侧内容区根据导航切换展示对应表单或记录列表")
add_bullet("完成后转向界面：无固定转向（内部标签页切换），底部快捷入口可跳转积分中心或成就中心")

add_section_label("B. 数据描述")
add_table("表6.1.7.1 个人中心-用户信息卡数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["头像", "VARCHAR(255)", "图片URL", "用户上传", "可编辑（点击头像触发上传组件）"],
    ["昵称", "VARCHAR(50)", "≤50字符", "注册/用户后修改", "可编辑"],
    ["年级", "VARCHAR(20)", "大一~研三", "注册/用户后修改", "可编辑"],
    ["专业", "VARCHAR(100)", "≤100字符", "注册/用户后修改", "可编辑"],
    ["学号", "VARCHAR(30)", "唯一", "注册时填写", "不可修改"],
    ["积分", "INT", "非负整数", "系统累计", "否（不可编辑）"],
    ["成就数", "INT", "非负整数", "系统累计", "否"],
    ["闯关数", "INT", "非负整数", "系统累计", "否"],
])
add_bullet("个人信息表单字段：昵称（≤50字符，非必填）、手机号（11位数字，选填校验格式）、邮箱（合法邮箱格式，选填校验格式）、年级（下拉选择，选填）、专业（≤100字符，选填）")
add_bullet("修改密码表单校验：原密码必填（需与数据库中BCrypt加密后的密码匹配）；新密码必填（≥6字符，不能与原密码相同）；确认密码必填（需与新密码一致）")
add_bullet("学习记录列表来源：调用GET /api/user/learning-records，数据格式[{type: 'browse'|'challenge', title, time, score?, duration?}]，分页加载")

add_section_label("C. 按键功能")
add_table("表6.1.7.2 个人中心按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["头像上传（点击头像区域）", "上传新头像", "调起文件选择器→限制图片格式jpg/png，≤2MB→上传至后端→更新头像URL"],
    ["导航菜单（个人信息/修改密码/学习记录）", "切换右侧内容区", "内部状态切换，无路由变化"],
    ["保存按钮（个人信息表单）", "保存个人信息修改", "校验字段格式→调用PUT /api/user/profile→成功后显示「保存成功」提示，更新左侧边栏的昵称/年级/专业"],
    ["保存按钮（修改密码表单）", "修改密码", "校验三字段非空→校验新密码与确认密码一致→调用PUT /api/user/password，Body{oldPassword, newPassword}→成功：提示「密码修改成功，请重新登录」，清除token并跳转登录页"],
    ["退出登录", "退出登录", "弹出确认对话框「确定退出登录？」→确定→调用POST /api/user/logout（后端清除Redis中的token）→清除前端localStorage中的token→跳转登录页"],
    ["积分中心/我的成就（快捷入口）", "跳转积分/成就页", "路由/score//achievement"],
])

# 6.1.8 智能客服模块
add_h3("6.1.8 智能客服模块")

add_h4("6.1.8.1 智能客服界面（VC801）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：学生、管理员")
add_bullet("主要功能：基于DeepSeek API的智能反诈助手对话界面。左侧为历史会话列表，右侧为聊天主区域。支持多轮对话、创建/删除会话、清空对话、快捷问题发送和反馈评价")
add_bullet("完成后转向界面：无固定转向")

add_section_label("B. 数据描述")
add_table("表6.1.8.1 智能客服数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["会话ID", "BIGINT", "自增主键", "创建新会话时自动生成", "否"],
    ["消息角色", "ENUM", "'user'/'assistant'", "系统自动标记", "否"],
    ["消息内容", "TEXT", "用户输入或AI回复原文", "用户输入触发AI生成", "否（只读历史）"],
    ["消息时间", "DATETIME", "yyyy-MM-dd HH:mm:ss", "自动", "否"],
    ["反馈状态", "TINYINT", "1=满意，-1=不满意，0=未评价", "用户点击反馈按钮", "可修改（再次点击切换）"],
])
add_bullet("数据存储：消息历史保存在MySQL chat_message表；会话列表按last_update_time DESC排序；清空对话=逻辑删除会话下所有消息（打标记保留在数据库）")

add_section_label("C. 按键功能")
add_table("表6.1.8.2 智能客服按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["新会话按钮（左侧）", "创建新会话", "调用POST /api/chat/session/create，成功后左侧新增一条会话记录，右侧切换为空会话欢迎页"],
    ["删除会话按钮", "删除历史会话", "弹出确认对话框「确定删除此会话？」→确定→调用DELETE /api/chat/session/{sessionId}，物理删除该会话及其所有消息"],
    ["会话卡片（左侧列表）", "切换会话", "调用GET /api/chat/session/{sessionId}/messages加载该会话的历史消息"],
    ["清空对话按钮", "清空当前会话消息", "弹出确认对话框「确定清空当前对话？」→确定→调用DELETE /api/chat/session/{sessionId}/messages，逻辑删除所有消息"],
    ["发送按钮/Ctrl+Enter", "发送消息", "输入框非空校验→将用户消息添加到界面（乐观更新）→调用POST /api/chat/ask，Body{sessionId, content}→后端调用DeepSeek API生成回复→返回后展示助手消息"],
    ["快捷问题标签（欢迎页）", "自动发送预设问题", "点击后自动填入输入框并触发发送，等同用户手动输入后发送"],
    ["满意/不满意按钮", "提交反馈", "调用POST /api/chat/feedback，Body{messageId, feedback: 1/-1}，按钮状态切换"],
])
add_bullet("发送消息后端逻辑：接收用户消息→携带当前会话历史拼接成prompt→调用DeepSeek API→将「用户消息+助手回复」批量保存到数据库→返回助手的回复给前端")

# 6.1.9 智能推荐模块
add_h3("6.1.9 智能推荐模块")

add_h4("6.1.9.1 智能推荐界面（VC901）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：已登录学生")
add_bullet("主要功能：基于用户画像和混合推荐算法（内容推荐TF-IDF+余弦相似度+协同过滤皮尔逊相关系数+KNN+SPM序列模式挖掘），为用户推荐案例、资讯和闯关内容。顶部展示用户当前生命周期阶段、知识水平分数和薄弱环节标签")
add_bullet("完成后转向界面：点击推荐卡片→跳转对应详情界面")

add_section_label("B. 数据描述")
add_table("表6.1.9.1 智能推荐数据描述", ["字段", "类型", "格式/说明", "生成方式", "是否允许编辑"], [
    ["生命周期阶段", "ENUM", "'newbie'(新手期)/'growing'(成长期)/'mature'(成熟期)", "后端基于用户行为自动计算", "否"],
    ["知识水平", "INT", "0~100分", "根据闯关成绩和浏览情况综合评估", "否"],
    ["薄弱环节标签", "JSON Array", "[\"冒充公检法\",\"投资理财\"]", "根据答题错误率最高的知识点提取", "否"],
    ["推荐项列表", "JSON Array", "[{itemType:'case'|'news'|'challenge', itemId, title, reason}]", "后端推荐算法计算", "否"],
])
add_bullet("推荐算法：newbie（冷启动）→内容推荐（TF-IDF+余弦相似度）；growing→混合推荐（内容推荐+协同过滤加权融合）；mature→协同过滤（皮尔逊相关系数+KNN+SPM序列模式挖掘）")
add_bullet("冷启动用户提示：当用户处于newbie阶段且行为数据不足时，在推荐列表顶部显示引导提示「建议先浏览一些案例和资讯，系统将为您优化推荐」")

add_section_label("C. 按键功能")
add_table("表6.1.9.2 智能推荐按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["标签页（案例推荐/资讯推荐/闯关推荐）", "切换推荐类型", "切换后请求对应的推荐列表，调用GET /api/recommend?type=case|news|challenge"],
    ["推荐卡片", "跳转详情", "根据itemType路由至对应详情：case→/case/:id，news→/news/:id，challenge→/challenge/:id"],
])

# 6.1.10 积分成就模块
add_h3("6.1.10 积分成就模块")

add_h4("6.1.10.1 积分中心界面（VC1001）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：已登录学生")
add_bullet("主要功能：展示用户当前等级徽章、总积分、本周积分、等级进度条、距离下一等级的所需积分。下方展示已解锁/总成就数统计、积分获取规则列表和等级特权预览")
add_bullet("完成后转向界面：点击成就统计卡片可跳转成就中心界面（VC1002）")

add_section_label("B. 数据描述")
add_table("表6.1.10.1 积分中心数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["当前等级", "TINYINT", "Lv1~Lv10", "积分到达对应门槛自动升级", "否"],
    ["等级名称", "VARCHAR(50)", "青铜守护者/白银卫士/黄金精英等", "系统预设映射", "否"],
    ["总积分", "INT", "非负整数", "浏览/闯关/发帖/评论/解锁成就累计", "否"],
    ["本周积分", "INT", "非负整数", "本周一00:00以来的积分获得总和", "否"],
    ["等级进度", "—", "当前积分/升级所需积分×100%", "自动计算", "否"],
])
add_bullet("积分规则：浏览案例+2（每日上限20）；完成答题挑战基础+5（不限）；S评级+20；情景模拟+10~15；发布帖子+3（每日上限5）；评论互动+1（每日上限10）；解锁成就+5~200")
add_bullet("等级门槛：Lv1(0)反诈新手→Lv2(50)反诈学徒→Lv3(200)青铜守护者→Lv4(500)白银卫士→Lv5(1000)黄金精英→Lv6(2000)铂金侦探→Lv7(3500)钻石专家→Lv8(5000)大师级→Lv9(7500)宗师级→Lv10(10000)反诈传奇")
add_bullet("等级特权预览：展示当前等级前后各两个等级的名称和特权描述")

add_section_label("C. 按键功能")
add_table("表6.1.10.2 积分中心按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["成就统计卡片（「已解锁X/Y」）", "跳转成就中心", "路由跳转至/achievement，可携带?tab=unlocked参数"],
])

add_h4("6.1.10.2 成就中心界面（VC1002）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：已登录学生")
add_bullet("主要功能：以卡片网格展示所有成就项，支持按全部/已解锁/未解锁过滤。每张卡片展示成就名称、图标、描述、积分奖励和解锁状态")
add_bullet("完成后转向界面：无固定转向")

add_section_label("B. 数据描述")
add_table("表6.1.10.3 成就中心数据描述", ["字段", "类型", "长度/格式", "生成方式", "是否允许编辑"], [
    ["成就ID", "BIGINT", "自增主键", "系统自动", "否"],
    ["名称", "VARCHAR(100)", "≤100字符", "管理员预设", "管理员可编辑"],
    ["描述", "VARCHAR(200)", "成就条件说明", "管理员预设", "管理员可编辑"],
    ["图标", "VARCHAR(50)", "Emoji字符", "管理员预设", "管理员可编辑"],
    ["积分奖励", "INT", "正整数", "管理员预设", "管理员可编辑"],
    ["解锁条件类型", "ENUM", "login_count/browse_count/challenge_count/perfect_score/post_count/continuous_days/challenge_complete", "管理员预设", "管理员可编辑"],
    ["条件值", "INT", "不同类型的阈值", "管理员预设", "管理员可编辑"],
    ["是否已解锁", "BOOLEAN", "true/false", "系统自动判断", "否"],
])
add_bullet("成就示例：初识反诈（首次登录+5）、勤学苦练（浏览10案例+10）、反诈大师（浏览100案例+50）、初露锋芒（完成5闯关+15）、满分学霸（任意闯关满分+30）、社交达人（发布10帖子+20）、持之以恒（连续登录7天+100）、全知全能（完成所有闯关+200）")
add_bullet("解锁判断逻辑：每次用户行为触发（登录/浏览/闯关/发帖/评论），后端检查是否满足成就条件，满足则自动解锁并发放积分")

add_section_label("C. 按键功能")
add_table("表6.1.10.4 成就中心按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["标签页（全部/已解锁/未解锁）", "过滤成就列表", "切换后仅显示对应状态成就，默认「全部」"],
])

# 6.1.11 后台管理模块
add_h3("6.1.11 后台管理模块")

add_h4("6.1.11.1 数据看板界面（VC1101）")
add_section_label("A. 界面整体功能描述")
add_bullet("操作者：管理员")
add_bullet("主要功能：平台运营数据的可视化仪表盘，包含六张关键指标统计卡片、访问量趋势折线图、诈骗类型分布饼图、各院系平均得分柱状图、24小时用户活跃度热力图、TOP案例排行榜。支持数据刷新和导出（Excel/PDF）")
add_bullet("完成后转向界面：无固定转向")

add_section_label("B. 数据描述")
add_table("表6.1.11.1 后台管理-统计卡片数据描述", ["指标", "来源API", "格式"], [
    ["总用户数", "GET /api/admin/statistics/total-users", "INT"],
    ["今日活跃用户", "GET /api/admin/statistics/dau", "INT"],
    ["今日浏览量", "GET /api/admin/statistics/daily-views", "INT"],
    ["案例总数", "GET /api/admin/statistics/total-cases", "INT"],
    ["平均测试得分", "GET /api/admin/statistics/avg-score", "DECIMAL(5,1)"],
    ["闯关完成数", "GET /api/admin/statistics/challenge-completions", "INT"],
])
add_table("表6.1.11.2 后台管理-图表数据描述", ["图表", "数据来源", "数据格式", "参数"], [
    ["访问量趋势折线图", "GET /api/admin/statistics/visit-trend", "[{date, views, activeUsers, newUsers}]", "days=7|14|30"],
    ["诈骗类型分布饼图", "GET /api/admin/statistics/case-type-distribution", "[{type, count}]", "—"],
    ["各院系平均得分柱状图", "GET /api/admin/statistics/major-scores", "[{major, avgScore}]", "—"],
    ["24小时活跃度热力图", "GET /api/admin/statistics/hourly-activity", "[{hour, activityLevel}]", "—"],
    ["TOP案例排行榜", "GET /api/admin/statistics/top-cases", "[{rank, title, type, views, likes, heatScore}]", "limit=10"],
])
add_bullet("排行榜数据格式：排名TINYINT(1~10)，案例标题VARCHAR(200)，类型TINYINT枚举值，浏览量INT千分位显示（如21,345），点赞数INT，热度得分DECIMAL(5,1)（0~100，算法=views×0.4+likes×0.3+comments×0.2+shares×0.1）")

add_section_label("C. 按键功能")
add_table("表6.1.11.3 后台管理按键功能", ["按钮", "功能说明", "详细逻辑"], [
    ["刷新数据", "刷新所有数据", "重新调用所有统计API，展示最新数据，同时显示加载动画"],
    ["导出Excel", "导出统计数据为Excel", "调用GET /api/admin/export/excel?type=statistics→后端生成.xlsx文件并返回→前端触发下载。文件包含：每日统计表（日期/活跃用户/浏览量/新增用户）、院系统计表（院系/平均得分/参与人数）"],
    ["导出PDF", "导出现状报告为PDF", "前端截图所有图表+表格→调用POST /api/admin/export/pdf上传图片→后端生成包含图表和表格的PDF报告→返回下载链接。报告包含：数据快照、趋势图、分布图、排行榜"],
    ["时间范围切换（7天/14天/30天）", "切换趋势图时间范围", "重新请求访问量趋势数据，图表动画过渡"],
    ["排行表格", "—", "只读展示。点击案例标题可跳转至对应案例详情页（新标签页打开）"],
    ["各图表区域", "—", "鼠标悬停显示具体数据点数值，支持Tooltip"],
])

# 保存
doc.save(OUTPUT_PATH)
print(f"第六章文档已生成：{OUTPUT_PATH}")
print(f"段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")
