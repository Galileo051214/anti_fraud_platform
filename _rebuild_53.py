#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修正《系统设计说明书》5.3节：
1. 从备份恢复原始文档
2. 删除错误的5.3节功能子节
3. 按正确顺序重新生成功能子节
"""

import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# 文件路径
BASE_DIR = Path(r"D:\Project\anti_fraud_platform")
BACKUP_FILE = BASE_DIR / "docx" / "《项目阶段二》评审" / "专业综合课程设计-系统设计说明书-江乐霖-23201317-v2-backup.docx"
TARGET_FILE = BASE_DIR / "docx" / "《项目阶段二》评审" / "专业综合课程设计-系统设计说明书-江乐霖-23201317.docx"
CODE_WIKI_PATH = BASE_DIR / "CODE_WIKI.md"

# 功能配置数据 - 按模块组织（同之前版本，这里省略重复定义）
MODULES_CONFIG = {
    "5.3.1": {
        "name": "用户与认证模块",
        "preserve_existing": ["5.3.1.1", "5.3.1.2"],  # 保留已有的功能子节
        "functions": [
            {"id": "5.3.1.3", "name": "用户信息查询", "api": "GET /api/user/info",
             "desc": "获取当前登录用户的详细信息，包括用户基本信息、积分、等级等。通过JWT Token识别用户身份，返回用户VO对象。",
             "classes": ["UserController", "UserService", "UserServiceImpl", "UserMapper", "UserVO"],
             "files": ["UserController.java", "UserService.java", "UserServiceImpl.java", "UserMapper.java", "UserMapper.xml", "UserVO.java", "User.java", "JwtAuthenticationFilter.java", "LoginUser.java"]},
            {"id": "5.3.1.4", "name": "用户资料更新", "api": "PUT /api/user/update",
             "desc": "更新当前用户的个人资料，包括昵称、头像、年级、专业等信息。需校验数据合法性，更新用户画像相关信息。",
             "classes": ["UserController", "UserService", "UserServiceImpl", "UserMapper", "UserProfileService"],
             "files": ["UserController.java", "UserService.java", "UserServiceImpl.java", "UserMapper.java", "UserMapper.xml", "UpdateUserDTO.java", "User.java", "UserProfileService.java", "ProfileServiceImpl.java"]},
            {"id": "5.3.1.5", "name": "密码修改", "api": "PUT /api/user/password",
             "desc": "用户修改密码功能。需要验证原密码正确性，使用BCrypt加密新密码后更新数据库。修改成功后可选择是否强制重新登录。",
             "classes": ["UserController", "UserService", "UserServiceImpl", "BCryptPasswordEncoder", "UserMapper"],
             "files": ["UserController.java", "UserService.java", "UserServiceImpl.java", "UserMapper.java", "UserMapper.xml", "PasswordDTO.java", "User.java", "SecurityConfig.java"]},
            {"id": "5.3.1.6", "name": "用户登出", "api": "POST /api/user/logout",
             "desc": "用户登出功能。解析请求头中的JWT Token，计算剩余有效期，将Token写入Redis黑名单，实现Token失效机制。确保已登出Token无法再次使用。",
             "classes": ["UserController", "UserService", "UserServiceImpl", "JwtUtils", "RedisCacheUtil", "CacheConstants"],
             "files": ["UserController.java", "UserService.java", "UserServiceImpl.java", "JwtUtils.java", "RedisCacheUtil.java", "CacheConstants.java", "JwtAuthenticationFilter.java"]},
            {"id": "5.3.1.7", "name": "用户管理", "api": "GET /api/user/list, GET/PUT /api/user/{id}, PUT /api/user/{id}/enable, /{id}/disable",
             "desc": "管理员用户管理功能。支持分页查询用户列表、查看用户详情、更新用户信息、启用/禁用用户账号等操作。需要ADMIN角色权限。",
             "classes": ["UserController", "UserService", "UserServiceImpl", "UserMapper", "UserVO", "UserListDTO"],
             "files": ["UserController.java", "UserService.java", "UserServiceImpl.java", "UserMapper.java", "UserMapper.xml", "UserVO.java", "UserListDTO.java", "User.java", "SecurityConfig.java"]}
        ]
    },
    "5.3.2": {
        "name": "资讯学习模块",
        "preserve_existing": [],
        "functions": [
            {"id": "5.3.2.1", "name": "资讯分页列表", "api": "GET /api/news/page",
             "desc": "分页查询资讯列表，支持按分类ID、资讯类型、关键词筛选。返回结果按置顶优先、发布时间倒序排列，包含点赞数和是否已点赞标记。",
             "classes": ["NewsController", "NewsService", "NewsServiceImpl", "NewsMapper", "NewsVO"],
             "files": ["NewsController.java", "NewsService.java", "NewsServiceImpl.java", "NewsMapper.java", "NewsMapper.xml", "NewsVO.java", "News.java", "NewsCategory.java", "PageQueryDTO.java"]},
            {"id": "5.3.2.2", "name": "资讯详情查看", "api": "GET /api/news/{id}",
             "desc": "获取资讯详细信息，包括标题、内容、作者、发布时间、浏览量、点赞数等。同时返回分类信息和相关资讯推荐。",
             "classes": ["NewsController", "NewsService", "NewsServiceImpl", "NewsMapper", "NewsVO"],
             "files": ["NewsController.java", "NewsService.java", "NewsServiceImpl.java", "NewsMapper.java", "NewsMapper.xml", "NewsVO.java", "News.java", "RedisCacheUtil.java"]},
            {"id": "5.3.2.3", "name": "资讯发布", "api": "POST /api/news, POST /api/news/{id}/publish",
             "desc": "创建资讯并发布。管理员创建资讯记录，设置分类、类型、标签等属性。发布后更新发布时间，触发缓存失效和积分奖励。",
             "classes": ["NewsController", "NewsService", "NewsServiceImpl", "NewsMapper", "NewsCategoryMapper", "CacheRefreshService"],
             "files": ["NewsController.java", "NewsService.java", "NewsServiceImpl.java", "NewsMapper.java", "NewsMapper.xml", "News.java", "NewsDTO.java", "NewsCategoryMapper.java", "CacheRefreshService.java"]},
            {"id": "5.3.2.4", "name": "资讯编辑", "api": "PUT /api/news/{id}",
             "desc": "更新资讯内容、分类、标签等信息。编辑后自动更新修改时间，触发相关缓存失效。需要管理员或作者权限。",
             "classes": ["NewsController", "NewsService", "NewsServiceImpl", "NewsMapper", "CacheRefreshService"],
             "files": ["NewsController.java", "NewsService.java", "NewsServiceImpl.java", "NewsMapper.java", "NewsMapper.xml", "News.java", "NewsDTO.java", "CacheRefreshService.java"]},
            {"id": "5.3.2.5", "name": "资讯删除", "api": "DELETE /api/news/{id}",
             "desc": "逻辑删除资讯，将deleted字段置为1。级联删除相关点赞记录和浏览记录。触发缓存失效，确保数据一致性。",
             "classes": ["NewsController", "NewsService", "NewsServiceImpl", "NewsMapper", "NewsLikeMapper", "NewsBrowseLogMapper", "CacheRefreshService"],
             "files": ["NewsController.java", "NewsService.java", "NewsServiceImpl.java", "NewsMapper.java", "NewsLikeMapper.java", "NewsBrowseLogMapper.java", "News.java", "CacheRefreshService.java"]},
            {"id": "5.3.2.6", "name": "资讯浏览记录", "api": "POST /api/news/{id}/view, /{id}/browse, GET /api/news/browse/history",
             "desc": "记录用户浏览资讯的行为，包括浏览次数增加和停留时长记录。支持查询用户的浏览历史，用于推荐系统和用户画像更新。",
             "classes": ["NewsController", "NewsService", "NewsServiceImpl", "NewsBrowseLogMapper", "UserProfileService", "AchievementService"],
             "files": ["NewsController.java", "NewsService.java", "NewsServiceImpl.java", "NewsBrowseLogMapper.java", "NewsBrowseLog.java", "UserProfileService.java", "AchievementService.java", "RedisCacheUtil.java"]},
            {"id": "5.3.2.7", "name": "资讯点赞", "api": "POST/DELETE /api/news/{id}/like",
             "desc": "用户对资讯进行点赞或取消点赞操作。采用幂等设计，点赞前校验是否已点赞。点赞成功后更新点赞数并触发积分奖励。",
             "classes": ["NewsController", "NewsService", "NewsServiceImpl", "NewsLikeMapper", "NewsMapper", "ScoreService"],
             "files": ["NewsController.java", "NewsService.java", "NewsServiceImpl.java", "NewsLikeMapper.java", "NewsLike.java", "NewsMapper.java", "News.java", "ScoreService.java"]}
        ]
    },
    "5.3.3": {
        "name": "案例展示模块",
        "preserve_existing": [],
        "functions": [
            {"id": "5.3.3.1", "name": "案例分页列表", "api": "GET /api/case/page",
             "desc": "分页查询案例列表，支持按标签ID、关键词筛选。返回结果按精选优先、Wilson分数、发布时间排序，包含点赞率统计和用户是否已点赞标记。",
             "classes": ["CaseController", "FraudCaseService", "FraudCaseServiceImpl", "FraudCaseMapper", "CaseVO"],
             "files": ["CaseController.java", "FraudCaseService.java", "FraudCaseServiceImpl.java", "FraudCaseMapper.java", "FraudCaseMapper.xml", "CaseVO.java", "FraudCase.java", "CaseTag.java"]},
            {"id": "5.3.3.2", "name": "案例详情查看", "api": "GET /api/case/{id}",
             "desc": "获取案例详细信息，包括标题、内容、风险等级、目标人群、标签、Wilson分数等。返回相关案例推荐和防范建议。",
             "classes": ["CaseController", "FraudCaseService", "FraudCaseServiceImpl", "FraudCaseMapper", "CaseVO"],
             "files": ["CaseController.java", "FraudCaseService.java", "FraudCaseServiceImpl.java", "FraudCaseMapper.java", "FraudCaseMapper.xml", "CaseVO.java", "FraudCase.java", "RedisCacheUtil.java"]},
            {"id": "5.3.3.3", "name": "案例发布", "api": "POST /api/case, POST /api/case/{id}/publish",
             "desc": "创建案例并发布。管理员创建案例记录，设置风险等级、目标人群、标签等属性。计算并存储Wilson分数用于排序。发布后触发缓存失效。",
             "classes": ["CaseController", "FraudCaseService", "FraudCaseServiceImpl", "FraudCaseMapper", "CaseTagRelationMapper", "RecommendationAlgorithmUtil"],
             "files": ["CaseController.java", "FraudCaseService.java", "FraudCaseServiceImpl.java", "FraudCaseMapper.java", "FraudCase.java", "CaseDTO.java", "CaseTagRelationMapper.java", "RecommendationAlgorithmUtil.java"]},
            {"id": "5.3.3.4", "name": "案例编辑", "api": "PUT /api/case/{id}",
             "desc": "更新案例内容、标签、风险等级等信息。编辑后重新计算Wilson分数，更新修改时间，触发相关缓存失效。需要管理员权限。",
             "classes": ["CaseController", "FraudCaseService", "FraudCaseServiceImpl", "FraudCaseMapper", "CaseTagRelationMapper", "CacheRefreshService"],
             "files": ["CaseController.java", "FraudCaseService.java", "FraudCaseServiceImpl.java", "FraudCaseMapper.java", "CaseDTO.java", "FraudCase.java", "CaseTagRelationMapper.java", "CacheRefreshService.java"]},
            {"id": "5.3.3.5", "name": "案例删除", "api": "DELETE /api/case/{id}",
             "desc": "逻辑删除案例，将deleted字段置为1。级联删除相关标签关联、点赞记录和浏览记录。触发缓存失效，确保推荐系统数据一致性。",
             "classes": ["CaseController", "FraudCaseService", "FraudCaseServiceImpl", "FraudCaseMapper", "CaseTagRelationMapper", "CaseLikeMapper", "CaseBrowseLogMapper"],
             "files": ["CaseController.java", "FraudCaseService.java", "FraudCaseServiceImpl.java", "FraudCaseMapper.java", "CaseTagRelationMapper.java", "CaseLikeMapper.java", "CaseBrowseLogMapper.java", "CacheRefreshService.java"]},
            {"id": "5.3.3.6", "name": "案例浏览记录", "api": "POST /api/case/{id}/browse, GET /api/case/browse/history",
             "desc": "记录用户浏览案例的行为，首次浏览奖励+2积分，更新日/周/总排行榜。记录停留时长用于知识水平评估。支持查询用户浏览历史。",
             "classes": ["CaseController", "FraudCaseService", "FraudCaseServiceImpl", "CaseBrowseLogMapper", "ScoreService", "LeaderboardService", "AchievementService", "ProfileService"],
             "files": ["CaseController.java", "FraudCaseService.java", "FraudCaseServiceImpl.java", "CaseBrowseLogMapper.java", "CaseBrowseLog.java", "ScoreService.java", "LeaderboardService.java", "AchievementService.java", "ProfileServiceImpl.java"]},
            {"id": "5.3.3.7", "name": "案例点赞", "api": "POST/DELETE /api/case/{id}/like",
             "desc": "用户对案例进行点赞或取消点赞操作。点赞成功后更新案例点赞数和点赞率，重新计算Wilson分数。采用幂等设计防止重复点赞。",
             "classes": ["CaseController", "FraudCaseService", "FraudCaseServiceImpl", "CaseLikeMapper", "FraudCaseMapper", "RecommendationAlgorithmUtil"],
             "files": ["CaseController.java", "FraudCaseService.java", "FraudCaseServiceImpl.java", "CaseLikeMapper.java", "CaseLike.java", "FraudCaseMapper.java", "FraudCase.java", "RecommendationAlgorithmUtil.java"]},
            {"id": "5.3.3.8", "name": "热点案例推荐", "api": "GET /api/case/hot, /api/case/wilson",
             "desc": "获取热点案例推荐列表，基于浏览量、点赞数、评论数计算热度分，或基于Wilson置信区间计算点赞率排序。支持Redis缓存加速。",
             "classes": ["CaseController", "FraudCaseService", "FraudCaseServiceImpl", "FraudCaseMapper", "RecommendationAlgorithmUtil", "RedisCacheUtil"],
             "files": ["CaseController.java", "FraudCaseService.java", "FraudCaseServiceImpl.java", "FraudCaseMapper.java", "FraudCaseMapper.xml", "RecommendationAlgorithmUtil.java", "RedisCacheUtil.java", "CacheConstants.java"]}
        ]
    },
    "5.3.4": {
        "name": "知识闯关与情景模拟模块",
        "preserve_existing": [],
        "functions": [
            {"id": "5.3.4.1", "name": "关卡列表查询", "api": "GET /api/challenge/list",
             "desc": "查询用户的关卡列表，包含解锁状态和通关状态。按levelOrder排序，返回每个关卡的类型、难度、奖励积分等信息。",
             "classes": ["ChallengeController", "ChallengeService", "ChallengeServiceImpl", "ChallengeMapper", "UserChallengeRecordMapper"],
             "files": ["ChallengeController.java", "ChallengeService.java", "ChallengeServiceImpl.java", "ChallengeMapper.java", "Challenge.java", "UserChallengeRecordMapper.java", "UserChallengeRecord.java", "ChallengeVO.java"]},
            {"id": "5.3.4.2", "name": "关卡详情查看", "api": "GET /api/challenge/{id}",
             "desc": "获取关卡详细信息，包括题目内容、选项、分值等。对于情景模拟类型，返回情景剧本结构。需验证关卡解锁状态。",
             "classes": ["ChallengeController", "ChallengeService", "ChallengeServiceImpl", "ChallengeMapper", "ChallengeVO"],
             "files": ["ChallengeController.java", "ChallengeService.java", "ChallengeServiceImpl.java", "ChallengeMapper.java", "Challenge.java", "ChallengeVO.java", "ChallengeContent.java", "ScenarioScript.java"]},
            {"id": "5.3.4.3", "name": "答题闯关提交", "api": "POST /api/challenge/submit",
             "desc": "提交答题结果并评分。系统比对用户答案与正确答案，计算得分。通关后奖励积分并更新排行榜，触发成就检查（闯关数、满分、通关）。",
             "classes": ["ChallengeController", "ChallengeService", "ChallengeServiceImpl", "ChallengeMapper", "UserChallengeRecordMapper", "ScoreService", "LeaderboardService", "AchievementService"],
             "files": ["ChallengeController.java", "ChallengeService.java", "ChallengeServiceImpl.java", "ChallengeMapper.java", "UserChallengeRecordMapper.java", "SubmitChallengeDTO.java", "ScoreService.java", "LeaderboardService.java", "AchievementService.java", "ProfileServiceImpl.java"]},
            {"id": "5.3.4.4", "name": "闯关记录查询", "api": "GET /api/challenge/records, /api/challenge/progress",
             "desc": "查询用户的闯关历史记录和总体进度。返回每次闯关的得分、答题详情、通关状态，以及已完成关卡数、总积分等统计信息。",
             "classes": ["ChallengeController", "ChallengeService", "ChallengeServiceImpl", "UserChallengeRecordMapper", "ChallengeMapper"],
             "files": ["ChallengeController.java", "ChallengeService.java", "ChallengeServiceImpl.java", "UserChallengeRecordMapper.java", "UserChallengeRecordMapper.xml", "UserChallengeRecord.java", "ChallengeMapper.java", "ChallengeRecordVO.java"]},
            {"id": "5.3.4.5", "name": "情景模拟开始", "api": "POST /api/scenario/start/{challengeId}",
             "desc": "开始情景模拟关卡。初始化情景进度记录，设置当前节点为起始节点。返回首个场景的对话或决策选项。",
             "classes": ["ScenarioController", "ScenarioService", "ScenarioServiceImpl", "ScenarioProgressMapper", "ChallengeMapper"],
             "files": ["ScenarioController.java", "ScenarioService.java", "ScenarioServiceImpl.java", "ScenarioProgressMapper.java", "ScenarioProgress.java", "ChallengeMapper.java", "Challenge.java", "ScenarioScript.java"]},
            {"id": "5.3.4.6", "name": "情景决策推进", "api": "POST /api/scenario/decision",
             "desc": "用户在情景模拟中做出决策，推进状态机。根据当前节点和选择找到对应的边，记录决策是否安全选择，更新当前节点。到达终点时计算最终得分。",
             "classes": ["ScenarioController", "ScenarioService", "ScenarioServiceImpl", "ScenarioProgressMapper", "ScoreService", "LeaderboardService"],
             "files": ["ScenarioController.java", "ScenarioService.java", "ScenarioServiceImpl.java", "ScenarioProgressMapper.java", "DecisionDTO.java", "ScenarioProgress.java", "ScoreService.java", "LeaderboardService.java", "ScenarioScript.java"]},
            {"id": "5.3.4.7", "name": "情景进度查询", "api": "GET /api/scenario/progress/{challengeId}",
             "desc": "查询用户在某个情景模拟关卡中的当前进度。返回当前节点、决策历史、已完成状态、最终得分等信息。",
             "classes": ["ScenarioController", "ScenarioService", "ScenarioServiceImpl", "ScenarioProgressMapper", "ChallengeMapper"],
             "files": ["ScenarioController.java", "ScenarioService.java", "ScenarioServiceImpl.java", "ScenarioProgressMapper.java", "ScenarioProgress.java", "ChallengeMapper.java", "Challenge.java", "ScenarioScript.java"]}
        ]
    },
    "5.3.5": {
        "name": "社区互动模块",
        "preserve_existing": [],
        "functions": [
            {"id": "5.3.5.1", "name": "帖子分页列表", "api": "GET /api/forum/post/page",
             "desc": "分页查询社区帖子列表，支持按帖子类型、排序方式（时间/点赞/评论）、关键词筛选。返回帖子基本信息、作者信息、点赞数、评论数。",
             "classes": ["ForumController", "ForumPostService", "ForumPostServiceImpl", "ForumPostMapper", "PostVO"],
             "files": ["ForumController.java", "ForumPostService.java", "ForumPostServiceImpl.java", "ForumPostMapper.java", "ForumPostMapper.xml", "PostVO.java", "ForumPost.java", "PageQueryDTO.java"]},
            {"id": "5.3.5.2", "name": "帖子详情查看", "api": "GET /api/forum/post/{postId}",
             "desc": "获取帖子详细信息，包括标题、内容、作者、发布时间、浏览量、点赞数、评论数等。浏览量自动+1。",
             "classes": ["ForumController", "ForumPostService", "ForumPostServiceImpl", "ForumPostMapper", "PostVO"],
             "files": ["ForumController.java", "ForumPostService.java", "ForumPostServiceImpl.java", "ForumPostMapper.java", "ForumPostMapper.xml", "PostVO.java", "ForumPost.java"]},
            {"id": "5.3.5.3", "name": "帖子发布", "api": "POST /api/forum/post",
             "desc": "用户发布新帖子，支持经验分享、问题求助、讨论三种类型。发布成功奖励+3积分，更新日/周/总排行榜。",
             "classes": ["ForumController", "ForumPostService", "ForumPostServiceImpl", "ForumPostMapper", "ScoreService", "LeaderboardService"],
             "files": ["ForumController.java", "ForumPostService.java", "ForumPostServiceImpl.java", "ForumPostMapper.java", "ForumPost.java", "PostDTO.java", "ScoreService.java", "LeaderboardService.java"]},
            {"id": "5.3.5.4", "name": "帖子编辑", "api": "PUT /api/forum/post/{postId}",
             "desc": "更新帖子内容和类型。编辑后更新修改时间。需要帖子作者或管理员权限。",
             "classes": ["ForumController", "ForumPostService", "ForumPostServiceImpl", "ForumPostMapper"],
             "files": ["ForumController.java", "ForumPostService.java", "ForumPostServiceImpl.java", "ForumPostMapper.java", "ForumPost.java", "PostDTO.java"]},
            {"id": "5.3.5.5", "name": "帖子删除", "api": "DELETE /api/forum/post/{postId}",
             "desc": "逻辑删除帖子，将deleted字段置为1。级联删除帖子点赞、评论及评论点赞。需要帖子作者或管理员权限。",
             "classes": ["ForumController", "ForumPostService", "ForumPostServiceImpl", "ForumPostMapper", "PostLikeMapper", "CommentMapper", "CommentLikeMapper"],
             "files": ["ForumController.java", "ForumPostService.java", "ForumPostServiceImpl.java", "ForumPostMapper.java", "PostLikeMapper.java", "CommentMapper.java", "CommentLikeMapper.java", "ForumPost.java"]},
            {"id": "5.3.5.6", "name": "帖子点赞", "api": "POST/DELETE /api/forum/post/{postId}/like",
             "desc": "用户对帖子进行点赞或取消点赞操作。采用幂等设计，点赞前校验是否已点赞。点赞成功后更新帖子点赞数。",
             "classes": ["ForumController", "ForumPostService", "ForumPostServiceImpl", "PostLikeMapper", "ForumPostMapper"],
             "files": ["ForumController.java", "ForumPostService.java", "ForumPostServiceImpl.java", "PostLikeMapper.java", "PostLike.java", "ForumPostMapper.java", "ForumPost.java"]},
            {"id": "5.3.5.7", "name": "评论查看", "api": "GET /api/comment/post/{postId}, GET /api/forum/post/{postId}/comments",
             "desc": "获取帖子的评论列表，递归构建树形结构。返回评论内容、作者、点赞数、回复数，以及是否已点赞、是否为作者标记。",
             "classes": ["ForumController", "CommentController", "ForumPostService", "ForumPostServiceImpl", "CommentMapper", "CommentVO"],
             "files": ["ForumController.java", "CommentController.java", "ForumPostService.java", "ForumPostServiceImpl.java", "CommentMapper.java", "CommentMapper.xml", "CommentVO.java", "Comment.java"]},
            {"id": "5.3.5.8", "name": "评论发布", "api": "POST /api/comment",
             "desc": "用户对帖子或评论进行回复。支持多级回复（通过parent_id关联）。评论成功后更新帖子的评论数统计。",
             "classes": ["CommentController", "CommentService", "CommentServiceImpl", "CommentMapper", "ForumPostMapper"],
             "files": ["CommentController.java", "CommentService.java", "CommentServiceImpl.java", "CommentMapper.java", "Comment.java", "CommentDTO.java", "ForumPostMapper.java", "ForumPost.java"]}
        ]
    },
    "5.3.6": {
        "name": "AI智能客服模块",
        "preserve_existing": [],
        "functions": [
            {"id": "5.3.6.1", "name": "智能问答", "api": "POST /api/chat/ask",
             "desc": "用户向AI智能客服提问。系统使用DeepSeek大模型，结合反诈专业提示词生成回答。支持会话上下文记忆，记录Token消耗。",
             "classes": ["ChatController", "QAConversationService", "QAConversationServiceImpl", "DeepSeekClient", "AntiFraudPromptTemplate", "QAConversationMapper"],
             "files": ["ChatController.java", "QAConversationService.java", "QAConversationServiceImpl.java", "DeepSeekClient.java", "AntiFraudPromptTemplate.java", "QAConversationMapper.java", "QAConversation.java", "ChatDTO.java", "ChatVO.java"]},
            {"id": "5.3.6.2", "name": "会话历史查询", "api": "GET /api/chat/history/{sessionId}",
             "desc": "查询某个会话的完整对话历史记录。返回用户问题和AI回答的列表，支持分页。",
             "classes": ["ChatController", "QAConversationService", "QAConversationServiceImpl", "QAConversationMapper"],
             "files": ["ChatController.java", "QAConversationService.java", "QAConversationServiceImpl.java", "QAConversationMapper.java", "QAConversationMapper.xml", "QAConversation.java"]},
            {"id": "5.3.6.3", "name": "会话列表查询", "api": "GET /api/chat/sessions",
             "desc": "查询用户的所有AI对话会话列表。返回会话ID、创建时间、最后消息预览等信息，按时间倒序排列。",
             "classes": ["ChatController", "QAConversationService", "QAConversationServiceImpl", "QAConversationMapper"],
             "files": ["ChatController.java", "QAConversationService.java", "QAConversationServiceImpl.java", "QAConversationMapper.java", "QAConversationMapper.xml", "QAConversation.java", "SessionVO.java"]},
            {"id": "5.3.6.4", "name": "反馈提交", "api": "POST /api/chat/feedback",
             "desc": "用户对AI回答进行反馈（满意/不满意）。反馈数据用于评估AI服务质量，优化提示词和模型参数。",
             "classes": ["ChatController", "QAConversationService", "QAConversationServiceImpl", "QAConversationMapper"],
             "files": ["ChatController.java", "QAConversationService.java", "QAConversationServiceImpl.java", "QAConversationMapper.java", "QAConversationMapper.xml", "QAConversation.java", "FeedbackDTO.java"]}
        ]
    },
    "5.3.7": {
        "name": "个性化推荐模块",
        "preserve_existing": [],
        "functions": [
            {"id": "5.3.7.1", "name": "自动策略推荐", "api": "GET /api/recommendation/list",
             "desc": "根据用户画像的生命周期阶段（新手/成长/成熟）自动选择合适的推荐策略。新手使用冷启动策略，成长期使用内容+协同过滤，成熟期使用协同过滤。",
             "classes": ["RecommendationController", "RecommendationService", "RecommendationServiceImpl", "ProfileService", "RecommendationAlgorithmUtil"],
             "files": ["RecommendationController.java", "RecommendationService.java", "RecommendationServiceImpl.java", "ProfileService.java", "ProfileServiceImpl.java", "RecommendationAlgorithmUtil.java", "RecommendationVO.java"]},
            {"id": "5.3.7.2", "name": "新手推荐", "api": "GET /api/recommendation/newbie",
             "desc": "为新手用户提供推荐。优先推荐必读资讯，再根据静态属性（年级、专业）匹配目标人群案例，最后补充热点内容。评分公式：Wilson*0.6 + hot*0.4。",
             "classes": ["RecommendationController", "RecommendationService", "RecommendationServiceImpl", "NewsMapper", "FraudCaseMapper", "ChallengeMapper", "RecommendationAlgorithmUtil"],
             "files": ["RecommendationController.java", "RecommendationService.java", "RecommendationServiceImpl.java", "NewsMapper.java", "FraudCaseMapper.java", "ChallengeMapper.java", "RecommendationAlgorithmUtil.java", "FraudCase.java"]},
            {"id": "5.3.7.3", "name": "成长期推荐", "api": "GET /api/recommendation/growing",
             "desc": "为成长期用户提供推荐。使用内容推荐（用户行为向量与案例标签向量余弦相似度）+ SPM序列预测 + 上下文修正（弱点标签×1.2、低知识水平×1.1）。综合分=0.5*content + 0.3*context + 0.2*hot。",
             "classes": ["RecommendationController", "RecommendationService", "RecommendationServiceImpl", "UserBehaviorMatrixMapper", "AssociationRuleMapper", "RecommendationAlgorithmUtil"],
             "files": ["RecommendationController.java", "RecommendationService.java", "RecommendationServiceImpl.java", "UserBehaviorMatrixMapper.java", "AssociationRuleMapper.java", "RecommendationAlgorithmUtil.java", "UserBehaviorMatrix.java", "AssociationRule.java"]},
            {"id": "5.3.7.4", "name": "成熟期推荐", "api": "GET /api/recommendation/mature",
             "desc": "为成熟期用户提供推荐。使用协同过滤策略：同年级预过滤 → KNN(k=20) → 皮尔逊相关系数计算相似度 → 偏置加权平均预测。将推荐标签转换为具体案例。",
             "classes": ["RecommendationController", "RecommendationService", "RecommendationServiceImpl", "UserSimilarityMapper", "FraudCaseMapper", "RecommendationAlgorithmUtil"],
             "files": ["RecommendationController.java", "RecommendationService.java", "RecommendationServiceImpl.java", "UserSimilarityMapper.java", "UserSimilarity.java", "FraudCaseMapper.java", "RecommendationAlgorithmUtil.java"]},
            {"id": "5.3.7.5", "name": "用户兴趣分析", "api": "GET /api/recommendation/interest",
             "desc": "分析用户的兴趣标签分布，基于浏览、点赞、闯关等行为数据计算各标签的兴趣分数，归一化到0-100并缓存。",
             "classes": ["RecommendationController", "RecommendationService", "RecommendationServiceImpl", "UserBehaviorMatrixMapper", "RedisCacheUtil"],
             "files": ["RecommendationController.java", "RecommendationService.java", "RecommendationServiceImpl.java", "UserBehaviorMatrixMapper.java", "UserBehaviorMatrix.java", "RedisCacheUtil.java", "CacheConstants.java"]}
        ]
    },
    "5.3.8": {
        "name": "数据统计模块",
        "preserve_existing": [],
        "functions": [
            {"id": "5.3.8.1", "name": "数据看板聚合", "api": "GET /api/statistics/dashboard",
             "desc": "聚合展示管理员数据看板，包括今日浏览量、新增用户、活跃用户、通关数、平均分等核心指标。每个子查询有try/catch兜底，确保看板稳定展示。",
             "classes": ["StatisticsController", "StatisticsService", "StatisticsServiceImpl", "StatisticsQueryMapper", "DailyStatisticsMapper"],
             "files": ["StatisticsController.java", "StatisticsService.java", "StatisticsServiceImpl.java", "StatisticsQueryMapper.java", "StatisticsQueryMapper.xml", "DailyStatisticsMapper.java", "DashboardVO.java"]},
            {"id": "5.3.8.2", "name": "访问趋势分析", "api": "GET /api/statistics/visit/trend",
             "desc": "查询访问趋势数据，按日期统计浏览量、活跃用户数、新增用户数等。支持指定时间范围，用于绘制趋势图表。",
             "classes": ["StatisticsController", "StatisticsService", "StatisticsServiceImpl", "StatisticsQueryMapper", "DailyStatisticsMapper"],
             "files": ["StatisticsController.java", "StatisticsService.java", "StatisticsServiceImpl.java", "StatisticsQueryMapper.java", "StatisticsQueryMapper.xml", "DailyStatisticsMapper.java", "DailyStatistics.java"]},
            {"id": "5.3.8.3", "name": "诈骗类型分布", "api": "GET /api/statistics/fraud/types",
             "desc": "统计各类诈骗案例的浏览量、点赞数、评论数分布。按案例标签聚合数据，用于分析学生关注的诈骗类型。",
             "classes": ["StatisticsController", "StatisticsService", "StatisticsServiceImpl", "StatisticsQueryMapper", "CaseTagMapper"],
             "files": ["StatisticsController.java", "StatisticsService.java", "StatisticsServiceImpl.java", "StatisticsQueryMapper.java", "StatisticsQueryMapper.xml", "CaseTagMapper.java", "CaseTag.java", "FraudCaseMapper.java"]},
            {"id": "5.3.8.4", "name": "院系积分统计", "api": "GET /api/statistics/department/scores",
             "desc": "按院系统计学生积分总和、平均积分、通关人数、发帖数等指标。支持排名查询，用于院系竞赛和激励。",
             "classes": ["StatisticsController", "StatisticsService", "StatisticsServiceImpl", "StatisticsQueryMapper", "DepartmentStatisticsMapper"],
             "files": ["StatisticsController.java", "StatisticsService.java", "StatisticsServiceImpl.java", "StatisticsQueryMapper.java", "StatisticsQueryMapper.xml", "DepartmentStatisticsMapper.java", "DepartmentStatistics.java", "UserMapper.java"]},
            {"id": "5.3.8.5", "name": "统计数据导出", "api": "GET /api/statistics/export/daily, /api/statistics/export/department",
             "desc": "使用EasyExcel导出统计数据。支持导出每日统计报表和院系统计报表，生成Excel文件供管理员下载分析。",
             "classes": ["StatisticsController", "StatisticsService", "StatisticsServiceImpl", "DailyStatisticsMapper", "DepartmentStatisticsMapper"],
             "files": ["StatisticsController.java", "StatisticsService.java", "StatisticsServiceImpl.java", "DailyStatisticsMapper.java", "DepartmentStatisticsMapper.java", "DailyStatistics.java", "DepartmentStatistics.java", "ExcelUtil.java"]}
        ]
    },
    "5.3.9": {
        "name": "积分成就与排行榜模块",
        "preserve_existing": [],
        "functions": [
            {"id": "5.3.9.1", "name": "积分信息查询", "api": "GET /api/score/info",
             "desc": "查询当前用户的积分信息，包括总积分、当前等级、周积分。等级计算公式：level = totalScore/100 + 1。支持查询其他用户积分（管理员权限）。",
             "classes": ["ScoreController", "ScoreService", "ScoreServiceImpl", "UserScoreMapper", "UserScore"],
             "files": ["ScoreController.java", "ScoreService.java", "ScoreServiceImpl.java", "UserScoreMapper.java", "UserScoreMapper.xml", "UserScore.java", "User.java", "LoginUser.java"]},
            {"id": "5.3.9.2", "name": "成就列表查询", "api": "GET /api/achievement/list",
             "desc": "查询系统所有成就列表，包括成就代码、名称、描述、条件类型、条件值、奖励积分。用于展示成就图鉴。",
             "classes": ["AchievementController", "AchievementService", "AchievementServiceImpl", "AchievementMapper", "Achievement"],
             "files": ["AchievementController.java", "AchievementService.java", "AchievementServiceImpl.java", "AchievementMapper.java", "AchievementMapper.xml", "Achievement.java", "AchievementVO.java"]},
            {"id": "5.3.9.3", "name": "用户成就查询", "api": "GET /api/achievement/user",
             "desc": "查询当前用户已解锁的成就列表。返回成就详情、解锁时间，以及已解锁/总成就数统计。支持分页查询。",
             "classes": ["AchievementController", "AchievementService", "AchievementServiceImpl", "UserAchievementMapper", "AchievementMapper"],
             "files": ["AchievementController.java", "AchievementService.java", "AchievementServiceImpl.java", "UserAchievementMapper.java", "UserAchievementMapper.xml", "UserAchievement.java", "AchievementMapper.java", "UserAchievementVO.java"]},
            {"id": "5.3.9.4", "name": "日榜查询", "api": "GET /api/leaderboard/daily",
             "desc": "查询每日积分排行榜，展示当天积分最高的用户列表。支持分页，返回排名、用户信息、积分。使用Redis缓存，TTL=300秒。",
             "classes": ["LeaderboardController", "LeaderboardService", "LeaderboardServiceImpl", "LeaderboardMapper", "RedisCacheUtil"],
             "files": ["LeaderboardController.java", "LeaderboardService.java", "LeaderboardServiceImpl.java", "LeaderboardMapper.java", "LeaderboardMapper.xml", "Leaderboard.java", "RedisCacheUtil.java", "CacheConstants.java"]},
            {"id": "5.3.9.5", "name": "周榜查询", "api": "GET /api/leaderboard/weekly",
             "desc": "查询每周积分排行榜，展示本周积分最高的用户列表。每日0点自动重排，使用Redis缓存，TTL=900秒。",
             "classes": ["LeaderboardController", "LeaderboardService", "LeaderboardServiceImpl", "LeaderboardMapper", "RedisCacheUtil"],
             "files": ["LeaderboardController.java", "LeaderboardService.java", "LeaderboardServiceImpl.java", "LeaderboardMapper.java", "LeaderboardMapper.xml", "Leaderboard.java", "RedisCacheUtil.java", "CacheConstants.java"]},
            {"id": "5.3.9.6", "name": "总榜查询", "api": "GET /api/leaderboard/all",
             "desc": "查询总积分排行榜，展示历史累计积分最高的用户列表。每日0点自动重排，使用Redis缓存，TTL=1800秒。",
             "classes": ["LeaderboardController", "LeaderboardService", "LeaderboardServiceImpl", "LeaderboardMapper", "RedisCacheUtil"],
             "files": ["LeaderboardController.java", "LeaderboardService.java", "LeaderboardServiceImpl.java", "LeaderboardMapper.java", "LeaderboardMapper.xml", "Leaderboard.java", "RedisCacheUtil.java", "CacheConstants.java"]},
            {"id": "5.3.9.7", "name": "用户排名查询", "api": "GET /api/leaderboard/user-rank",
             "desc": "查询当前用户在日榜、周榜、总榜中的排名位置。无需分页，直接返回三个榜的排名数据。",
             "classes": ["LeaderboardController", "LeaderboardService", "LeaderboardServiceImpl", "LeaderboardMapper", "UserScoreMapper"],
             "files": ["LeaderboardController.java", "LeaderboardService.java", "LeaderboardServiceImpl.java", "LeaderboardMapper.java", "LeaderboardMapper.xml", "UserScoreMapper.java", "UserScore.java", "UserRankVO.java"]}
        ]
    }
}


def insert_paragraph_after(doc, paragraph, new_paragraph):
    """在指定段落之后插入新段落"""
    new_p = new_paragraph._element
    paragraph._element.addnext(new_p)


def create_heading4(doc, text):
    """创建Heading 4标题段落"""
    para = doc.add_paragraph(text, style='Heading 4')
    return para


def create_normal_paragraph(doc, text):
    """创建Normal段落"""
    para = doc.add_paragraph(text, style='Normal')
    return para


def create_table(doc, func_config):
    """创建文件列表表格"""
    files = func_config['files'][:8]  # 最多8个文件
    table = doc.add_table(rows=len(files)+1, cols=4)
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    headers = ['名称', '类型', '存放位置', '说明']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    
    # 填充文件数据
    for i, file_name in enumerate(files, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = file_name
        row_cells[1].text = 'Java类' if file_name.endswith('.java') else '配置文件'
        row_cells[2].text = f'com.anti.{file_name.replace(".java", "").lower()}'
        row_cells[3].text = f'{func_config["name"]}相关'
    
    return table


def create_function_elements(doc, func_config):
    """创建功能子节的所有元素（段落+表格）"""
    elements = []
    
    # 1. Heading 4 标题
    heading_para = doc.add_paragraph(func_config['name'], style='Heading 4')
    elements.append(('paragraph', heading_para))
    
    # 2. 功能设计描述段落
    desc_para = doc.add_paragraph(style='Normal')
    desc_para.add_run('1 功能设计描述').bold = True
    desc_para.add_run(f'\n{func_config["desc"]}')
    desc_para.add_run(f'\nAPI路径：{func_config["api"]}')
    elements.append(('paragraph', desc_para))
    
    # 3. 类段落
    class_para = doc.add_paragraph(style='Normal')
    class_para.add_run('（1）类').bold = True
    classes_text = '、'.join(func_config['classes'][:5])
    class_para.add_run(f'\n{classes_text}。')
    elements.append(('paragraph', class_para))
    
    # 4. 文件列表段落
    file_para = doc.add_paragraph(style='Normal')
    file_para.add_run('（3）文件列表').bold = True
    file_para.add_run(f'如表{func_config["id"].replace(".", "-")}所示。')
    elements.append(('paragraph', file_para))
    
    # 5. 表格标题段落（居中）
    table_caption = doc.add_paragraph(f'表{func_config["id"].replace(".", "-")}  {func_config["name"]}文件列表', style='Normal')
    table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elements.append(('paragraph', table_caption))
    
    # 6. 表格
    table = create_table(doc, func_config)
    elements.append(('table', table))
    
    # 7. 空段落
    empty_para = doc.add_paragraph(style='Normal')
    elements.append(('paragraph', empty_para))
    
    return elements


def find_module_heading_index(doc, module_id):
    """查找模块标题的索引"""
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 3' and module_id in para.text:
            return i
    return -1


def find_last_preserved_function_index(doc, module_idx, preserved_ids):
    """查找需要保留的最后一个功能子节的结束位置"""
    if not preserved_ids:
        return module_idx
    
    last_idx = module_idx
    for i in range(module_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style.name == 'Heading 3':
            # 到达下一个模块，停止
            break
        if para.style.name == 'Heading 4':
            # 检查是否是需要保留的功能子节
            for preserved_id in preserved_ids:
                if preserved_id in para.text:
                    last_idx = i
                    # 找到这个功能子节的内容结束位置（下一个Heading或模块结束）
                    for j in range(i + 1, len(doc.paragraphs)):
                        if doc.paragraphs[j].style.name in ['Heading 3', 'Heading 4']:
                            last_idx = j - 1
                            break
                    break
    
    return last_idx


def delete_heading4_sections(doc, module_idx, preserved_ids):
    """删除模块内的Heading 4功能子节（除了需要保留的）"""
    elements_to_delete = []
    
    # 找到模块结束位置（下一个Heading 3或Heading 2）
    module_end_idx = len(doc.paragraphs)
    for i in range(module_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name in ['Heading 2', 'Heading 3']:
            module_end_idx = i
            break
    
    # 遍历模块内的所有段落
    current_func_start = -1
    is_preserved_func = False
    
    for i in range(module_idx + 1, module_end_idx):
        para = doc.paragraphs[i]
        
        if para.style.name == 'Heading 4':
            # 新的功能子节开始
            if current_func_start != -1 and not is_preserved_func:
                # 上一个功能子节需要删除
                for j in range(current_func_start, i):
                    elements_to_delete.append(j)
            
            current_func_start = i
            # 检查是否需要保留
            is_preserved_func = False
            for preserved_id in preserved_ids:
                if preserved_id in para.text:
                    is_preserved_func = True
                    break
        
        # 处理最后一个功能子节
        if i == module_end_idx - 1 and current_func_start != -1 and not is_preserved_func:
            # 到达模块结束，删除最后一个功能子节
            # 需要删除从current_func_start到模块结束前所有内容
            # 但要保留模块简介段落（在模块标题后）
            # 先找到功能子节的实际结束位置
            for j in range(current_func_start, module_end_idx):
                elements_to_delete.append(j)
    
    # 删除元素（从后向前）
    for idx in sorted(elements_to_delete, reverse=True):
        para = doc.paragraphs[idx]
        para._element.getparent().remove(para._element)
    
    # 同时删除相关的表格（功能子节内的表格）
    # 简化处理：删除所有在模块范围内的表格（保留表格数量需要更精确的逻辑）
    
    return len(elements_to_delete)


def process_module(doc, module_id, module_config):
    """处理一个模块"""
    print(f"\n处理模块 {module_id} {module_config['name']}...")
    
    # 查找模块标题位置
    module_idx = find_module_heading_index(doc, module_id)
    if module_idx == -1:
        print(f"  警告：未找到模块标题 {module_id}")
        return
    
    print(f"  模块标题位置：{module_idx}")
    
    # 删除不需要保留的功能子节
    preserved_ids = module_config.get('preserve_existing', [])
    deleted_count = delete_heading4_sections(doc, module_idx, preserved_ids)
    print(f"  删除了 {deleted_count} 个段落")
    
    # 重新查找模块标题位置（因为段落已删除）
    module_idx = find_module_heading_index(doc, module_id)
    
    # 找到插入位置
    if preserved_ids:
        # 找到最后一个保留的功能子节的位置
        insert_after_idx = find_last_preserved_function_index(doc, module_idx, preserved_ids)
        print(f"  在保留功能子节之后插入，位置：{insert_after_idx}")
        # 只插入不需要保留的功能
        funcs_to_insert = [f for f in module_config['functions'] 
                          if f['id'] not in preserved_ids and f['id'] > max(preserved_ids)]
    else:
        # 找到模块简介段落结束位置（通常是模块标题后的2-3个Normal段落）
        insert_after_idx = module_idx
        for i in range(module_idx + 1, min(module_idx + 5, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if para.style.name == 'Normal' and para.text.strip():
                insert_after_idx = i
            elif para.style.name.startswith('Heading'):
                break
        
        print(f"  在模块简介之后插入，位置：{insert_after_idx}")
        funcs_to_insert = module_config['functions']
    
    # 直接在主文档的末尾添加新功能子节
    # python-docx会自动添加到文档末尾，我们需要在正确的位置插入
    
    # 方案：先创建所有段落，然后一次性移动到正确位置
    new_elements = []
    for func in funcs_to_insert:
        print(f"  准备插入：{func['id']} {func['name']}")
        # 创建段落并记录其XML元素
        heading = doc.add_paragraph(func['name'], style='Heading 4')
        new_elements.append(('paragraph', heading._element))
        
        # 功能设计描述段落
        desc_para = doc.add_paragraph(style='Normal')
        desc_para.add_run('1 功能设计描述').bold = True
        desc_para.add_run(f'\n{func["desc"]}')
        desc_para.add_run(f'\nAPI路径：{func["api"]}')
        new_elements.append(('paragraph', desc_para._element))
        
        # 类段落
        class_para = doc.add_paragraph(style='Normal')
        class_para.add_run('（1）类').bold = True
        classes_text = '、'.join(func['classes'][:5])
        class_para.add_run(f'\n{classes_text}。')
        new_elements.append(('paragraph', class_para._element))
        
        # 文件列表段落
        file_para = doc.add_paragraph(style='Normal')
        file_para.add_run('（3）文件列表').bold = True
        file_para.add_run(f'如表{func["id"].replace(".", "-")}所示。')
        new_elements.append(('paragraph', file_para._element))
        
        # 表格标题段落（居中）
        table_caption = doc.add_paragraph(f'表{func["id"].replace(".", "-")}  {func["name"]}文件列表', style='Normal')
        table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_elements.append(('paragraph', table_caption._element))
        
        # 创建表格
        files = func['files'][:8]
        table = doc.add_table(rows=len(files)+1, cols=4)
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        headers = ['名称', '类型', '存放位置', '说明']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
        
        # 填充文件数据
        for i, file_name in enumerate(files, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = file_name
            row_cells[1].text = 'Java类' if file_name.endswith('.java') else '配置文件'
            row_cells[2].text = f'com.anti.{file_name.replace(".java", "").lower()}'
            row_cells[3].text = f'{func["name"]}相关'
        
        new_elements.append(('table', table._element))
        
        # 空段落
        empty_para = doc.add_paragraph(style='Normal')
        new_elements.append(('paragraph', empty_para._element))
    
    # 将新元素移动到正确位置
    insert_para = doc.paragraphs[insert_after_idx]
    for element_type, element in new_elements:
        # 从当前位置删除
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
        # 插入到正确位置
        insert_para._element.addnext(element)


def main():
    """主函数"""
    print("=" * 60)
    print("修正《系统设计说明书》5.3节")
    print("=" * 60)
    
    # 1. 从备份恢复
    print(f"\n步骤1：从备份恢复原始文档")
    if not BACKUP_FILE.exists():
        print(f"  错误：备份文件不存在！")
        return
    
    shutil.copy(BACKUP_FILE, TARGET_FILE)
    print(f"  已恢复备份文件")
    
    # 2. 打开文档
    print(f"\n步骤2：打开文档")
    doc = Document(TARGET_FILE)
    print(f"  段落数：{len(doc.paragraphs)}")
    print(f"  表格数：{len(doc.tables)}")
    
    # 3. 处理每个模块
    print(f"\n步骤3：处理各模块功能子节")
    
    for module_id in sorted(MODULES_CONFIG.keys()):
        module_config = MODULES_CONFIG[module_id]
        process_module(doc, module_id, module_config)
    
    # 4. 保存文档
    print(f"\n步骤4：保存文档")
    doc.save(TARGET_FILE)
    print(f"  已保存到：{TARGET_FILE}")
    
    # 5. 验证结果
    print(f"\n步骤5：验证结果")
    doc = Document(TARGET_FILE)
    
    # 统计各模块的功能子节数
    for module_id, module_config in MODULES_CONFIG.items():
        count = 0
        for para in doc.paragraphs:
            if para.style.name == 'Heading 4':
                for func in module_config['functions']:
                    if func['name'] in para.text or func['id'] in para.text:
                        count += 1
                        break
        expected = len(module_config['functions'])
        print(f"  {module_id} {module_config['name']}：{count} 个功能子节 (预期：{expected})")
    
    print("\n" + "=" * 60)
    print("修正完成！")
    print("=" * 60)


if __name__ == '__main__':
    main()