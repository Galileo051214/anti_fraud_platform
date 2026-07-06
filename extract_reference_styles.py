# -*- coding: utf-8 -*-
"""
提取参考文档 [2024]3-系统设计.docx 中第六章的样式信息
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import json

REF_PATH = r"D:\Project\anti_fraud_platform\docx\《项目阶段二》评审\[2024]3-系统设计.docx"
doc = Document(REF_PATH)

styles_info = {
    'headings': {},
    'paragraphs': {},
    'tables': []
}

# 遍历所有段落，提取第六章的内容和样式
in_chapter6 = False
chapter6_paragraphs = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith('6 用户界面设计概述'):
        in_chapter6 = True
    if in_chapter6 and text.startswith('7 '):
        break
    if in_chapter6:
        chapter6_paragraphs.append({
            'index': i,
            'text': text,
            'style_name': p.style.name,
            'alignment': p.alignment,
            'first_line_indent': p.paragraph_format.first_line_indent,
            'line_spacing': p.paragraph_format.line_spacing,
            'font_name': p.runs[0].font.name if p.runs else None,
            'font_size': p.runs[0].font.size if p.runs else None,
            'font_bold': p.runs[0].font.bold if p.runs else None,
            'font_color': p.runs[0].font.color.rgb if (p.runs and p.runs[0].font.color and p.runs[0].font.color.rgb) else None,
            'eastAsia_font': None
        })
        if p.runs and p.runs[0]._element.rPr:
            rFonts = p.runs[0]._element.rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                chapter6_paragraphs[-1]['eastAsia_font'] = rFonts.get(qn('w:eastAsia'))

# 提取第六章的表格
in_chapter6 = False
chapter6_tables = []

for i, table in enumerate(doc.tables):
    # 检查表格前的段落是否在第六章范围内
    prev_paragraph_idx = -1
    for j, p in enumerate(doc.paragraphs):
        if p._element is table._element.getprevious():
            prev_paragraph_idx = j
            break
    
    # 判断是否在第六章范围内
    chapter_start = None
    chapter_end = None
    for j, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith('6 用户界面设计概述'):
            chapter_start = j
        if text.startswith('7 '):
            chapter_end = j
            break
    
    if chapter_start is not None and (chapter_end is None or prev_paragraph_idx < chapter_end):
        if prev_paragraph_idx >= chapter_start:
            # 提取表格样式信息
            table_info = {
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'style': table.style.name,
                'header_font': {},
                'body_font': {}
            }
            
            # 提取表头样式
            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            table_info['header_font'] = {
                                'name': run.font.name,
                                'size': run.font.size,
                                'bold': run.font.bold,
                                'alignment': paragraph.alignment
                            }
                            break
                        break
                    break
            
            # 提取表体样式
            if len(table.rows) > 1:
                for cell in table.rows[1].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            table_info['body_font'] = {
                                'name': run.font.name,
                                'size': run.font.size,
                                'bold': run.font.bold
                            }
                            break
                        break
                    break
            
            chapter6_tables.append(table_info)

# 整理样式信息
print("=" * 60)
print("第六章段落样式详细信息")
print("=" * 60)
for p in chapter6_paragraphs:
    print(f"\n段落[{p['index']}]: {p['text'][:50]}...")
    print(f"  样式名: {p['style_name']}")
    print(f"  字体: {p['font_name']} / {p['eastAsia_font']}")
    print(f"  字号: {p['font_size']}")
    print(f"  加粗: {p['font_bold']}")
    print(f"  颜色: {p['font_color']}")
    print(f"  对齐: {p['alignment']}")
    print(f"  首行缩进: {p['first_line_indent']}")
    print(f"  行距: {p['line_spacing']}")

print("\n" + "=" * 60)
print("第六章表格样式详细信息")
print("=" * 60)
for t in chapter6_tables:
    print(f"\n表格[{t['index']}]: {t['rows']}行 x {t['cols']}列")
    print(f"  样式: {t['style']}")
    print(f"  表头字体: {t['header_font']}")
    print(f"  表体字体: {t['body_font']}")

# 保存样式信息到文件
style_data = {
    'paragraphs': chapter6_paragraphs,
    'tables': chapter6_tables
}
with open('reference_styles.json', 'w', encoding='utf-8') as f:
    json.dump(style_data, f, ensure_ascii=False, indent=2)

print("\n样式信息已保存到 reference_styles.json")
