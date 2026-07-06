# -*- coding: utf-8 -*-
"""
精确提取参考文档中各样式的中文字号
"""
from docx import Document
from docx.oxml.ns import qn

REF_PATH = r"D:\Project\anti_fraud_platform\docx\《项目阶段二》评审\[2024]3-系统设计.docx"
doc = Document(REF_PATH)

# EMU转pt: 1pt = 12700 EMU
def emu_to_pt(emu):
    if emu is None:
        return None
    return emu / 12700

# 检查第六章中实际段落的字体
in_chapter6 = False
samples = {}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith('6 用户界面设计概述'):
        in_chapter6 = True
    if in_chapter6 and text.startswith('7 '):
        break
    
    if in_chapter6 and text and p.runs:
        style_name = p.style.name
        if style_name not in samples:
            run = p.runs[0]
            font_size_emu = run.font.size
            font_size_pt = emu_to_pt(font_size_emu) if font_size_emu else None
            
            # 也获取样式定义的字号
            style_size_emu = None
            try:
                style_size_emu = p.style.font.size
            except:
                pass
            style_size_pt = emu_to_pt(style_size_emu) if style_size_emu else None
            
            # 中文字号对应
            cn_size = None
            if font_size_pt:
                if abs(font_size_pt - 16) < 0.5:
                    cn_size = '三号'
                elif abs(font_size_pt - 14) < 0.5:
                    cn_size = '四号'
                elif abs(font_size_pt - 12) < 0.5:
                    cn_size = '小四'
                elif abs(font_size_pt - 10.5) < 0.5:
                    cn_size = '五号'
                elif abs(font_size_pt - 9) < 0.5:
                    cn_size = '小五'
                else:
                    cn_size = f'{font_size_pt:.1f}pt'
            
            samples[style_name] = {
                'text': text[:30],
                'run_font_size_pt': font_size_pt,
                'style_font_size_pt': style_size_pt,
                'cn_size': cn_size,
                'bold': run.font.bold,
                'eastAsia': None
            }
            
            # 获取中文字体
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    samples[style_name]['eastAsia'] = rFonts.get(qn('w:eastAsia'))

print("参考文档第六章各样式字号:")
print("=" * 70)
for style, info in samples.items():
    print(f"\n【{style}】")
    print(f"  样本文本: {info['text']}")
    print(f"  字号: {info['run_font_size_pt']}pt ({info['cn_size']})")
    print(f"  样式定义字号: {info['style_font_size_pt']}pt")
    print(f"  加粗: {info['bold']}")
    print(f"  中文字体: {info['eastAsia']}")

# 检查表格字号
print("\n" + "=" * 70)
print("表格字号检查:")
print("=" * 70)

# 找到第六章的第一个表格
table_found = False
for i, table in enumerate(doc.tables):
    if table_found:
        break
    # 检查表头
    if len(table.rows) > 0:
        cell = table.rows[0].cells[0]
        if cell.paragraphs and cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            fs = emu_to_pt(run.font.size) if run.font.size else None
            cn = ''
            if fs:
                if abs(fs - 16) < 0.5: cn = '三号'
                elif abs(fs - 14) < 0.5: cn = '四号'
                elif abs(fs - 12) < 0.5: cn = '小四'
                elif abs(fs - 10.5) < 0.5: cn = '五号'
                elif abs(fs - 9) < 0.5: cn = '小五'
                else: cn = f'{fs:.1f}pt'
            print(f"表格{i}表头: {fs}pt ({cn}), 加粗={run.font.bold}")
    if len(table.rows) > 1:
        cell = table.rows[1].cells[0]
        if cell.paragraphs and cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            fs = emu_to_pt(run.font.size) if run.font.size else None
            cn = ''
            if fs:
                if abs(fs - 16) < 0.5: cn = '三号'
                elif abs(fs - 14) < 0.5: cn = '四号'
                elif abs(fs - 12) < 0.5: cn = '小四'
                elif abs(fs - 10.5) < 0.5: cn = '五号'
                elif abs(fs - 9) < 0.5: cn = '小五'
                else: cn = f'{fs:.1f}pt'
            print(f"表格{i}表体: {fs}pt ({cn}), 加粗={run.font.bold}")
            table_found = True
