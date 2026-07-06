# -*- coding: utf-8 -*-
from docx import Document

doc = Document(r'D:\Project\anti_fraud_platform\docx\《项目阶段二》评审\第六章-用户界面设计概述.docx')

def get_cn(pt):
    if pt is None: return 'None'
    if abs(pt-16)<0.5: return '三号'
    if abs(pt-14)<0.5: return '四号'
    if abs(pt-12)<0.5: return '小四'
    if abs(pt-10.5)<0.5: return '五号'
    return f'{pt:.1f}pt'

print('样式验证:')
for p in doc.paragraphs:
    s = p.style.name
    if not p.text.strip() or not p.runs: continue
    if s.startswith('Heading'):
        fs = p.runs[0].font.size/12700
        print(f'  {s}: {fs:.1f}pt ({get_cn(fs)}), 加粗={p.runs[0].font.bold}, 文本="{p.text}"')

# 表格
t = doc.tables[0]
fh = t.rows[0].cells[0].paragraphs[0].runs[0].font.size/12700
fb = t.rows[1].cells[0].paragraphs[0].runs[0].font.size/12700
print(f'\n表格表头: {fh:.1f}pt ({get_cn(fh)})')
print(f'表格表体: {fb:.1f}pt ({get_cn(fb)})')

print(f'\n段落总数: {len(doc.paragraphs)}')
print(f'表格总数: {len(doc.tables)}')

print('\n内容结构:')
for i, p in enumerate(doc.paragraphs):
    s = p.style.name
    if s.startswith('Heading'):
        print(f'  [{s}] {p.text}')
