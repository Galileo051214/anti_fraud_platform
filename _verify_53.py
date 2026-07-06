#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证《系统设计说明书》5.3节修正结果
"""

from pathlib import Path
from docx import Document

# 文件路径
BASE_DIR = Path(r"D:\Project\anti_fraud_platform")
TARGET_FILE = BASE_DIR / "docx" / "《项目阶段二》评审" / "专业综合课程设计-系统设计说明书-江乐霖-23201317.docx"

def main():
    """验证修正结果"""
    print("=" * 60)
    print("验证《系统设计说明书》5.3节修正结果")
    print("=" * 60)
    
    doc = Document(TARGET_FILE)
    print(f"\n文档统计：")
    print(f"  总段落数：{len(doc.paragraphs)}")
    print(f"  总表格数：{len(doc.tables)}")
    
    # 查找所有5.3节模块和功能子节
    print(f"\n5.3节结构分析：")
    
    # 按模块分组统计
    modules = {
        "5.3.1": {"name": "用户与认证模块", "expected": 7, "functions": []},
        "5.3.2": {"name": "资讯学习模块", "expected": 7, "functions": []},
        "5.3.3": {"name": "案例展示模块", "expected": 8, "functions": []},
        "5.3.4": {"name": "知识闯关与情景模拟模块", "expected": 7, "functions": []},
        "5.3.5": {"name": "社区互动模块", "expected": 8, "functions": []},
        "5.3.6": {"name": "AI智能客服模块", "expected": 4, "functions": []},
        "5.3.7": {"name": "个性化推荐模块", "expected": 5, "functions": []},
        "5.3.8": {"name": "数据统计模块", "expected": 5, "functions": []},
        "5.3.9": {"name": "积分成就与排行榜模块", "expected": 7, "functions": []}
    }
    
    # 统计每个模块的功能子节
    for para in doc.paragraphs:
        if para.style.name == 'Heading 3':
            # 模块标题
            for module_id in modules:
                if module_id in para.text:
                    print(f"\n模块标题：{para.text}")
                    break
        elif para.style.name == 'Heading 4':
            # 功能子节标题
            for module_id, module_info in modules.items():
                if para.text.startswith(module_id) or any(func_id in para.text for func_id in [f"{module_id}.{i}" for i in range(1, 10)]):
                    module_info['functions'].append(para.text)
                    break
    
    # 输出统计结果
    print(f"\n统计结果：")
    all_correct = True
    for module_id, module_info in modules.items():
        count = len(module_info['functions'])
        expected = module_info['expected']
        status = "✓" if count == expected else "✗"
        if count != expected:
            all_correct = False
        print(f"  {status} {module_id} {module_info['name']}: {count}个功能子节 (预期：{expected})")
        
        # 显示功能子节列表（按顺序）
        if module_info['functions']:
            print(f"    功能列表：")
            for i, func_name in enumerate(module_info['functions'], 1):
                # 检查顺序是否正确
                expected_id = f"{module_id}.{i}"
                if expected_id in func_name:
                    print(f"      ✓ {func_name}")
                else:
                    print(f"      ✗ {func_name} (预期：{expected_id})")
                    all_correct = False
    
    # 统计表格数量
    print(f"\n表格统计：")
    print(f"  总表格数：{len(doc.tables)}")
    
    # 预期表格数：每个功能子节1个表格
    expected_tables = sum(m['expected'] for m in modules.values())
    status = "✓" if len(doc.tables) >= expected_tables else "✗"
    print(f"  {status} 预期表格数：{expected_tables}")
    
    # 最终结论
    print(f"\n" + "=" * 60)
    if all_correct:
        print("✓ 修正成功！所有模块功能子节顺序和数量正确。")
    else:
        print("✗ 存在问题，需要进一步修正。")
    print("=" * 60)


if __name__ == '__main__':
    main()