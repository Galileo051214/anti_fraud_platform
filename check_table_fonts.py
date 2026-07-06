# -*- coding: utf-8 -*-
"""
精确提取第六章中表格的字号
"""
from docx import Document
from docx.oxml.ns import qn

REF_PATH = r"D:\Project\anti_fraud_platform\docx\《项目阶段二》评审\[2024]3-系统设计.docx"
doc = Document(REF_PATH)

def emu_to_pt(emu):
    if emu is None:
        return None
    return emu / 12700

def get_cn_size(pt):
    if pt is None:
        return 'None'
    if abs(pt - 16) < 0.5: return '三号'
    elif abs(pt - 14) < 0.5: return '四号'
    elif abs(pt - 12) < 0.5: return '小四'
    elif abs(pt - 10.5) < 0.5: return '五号'
    elif abs(pt - 9) < 0.5: return '小五'
    else: return f'{pt:.1f}pt'

# 找到第六章起始段落索引
chapter6_start = None
chapter6_end = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith('6 用户界面设计概述'):
        chapter6_start = i
    if chapter6_start is not None and text.startswith('7 '):
        chapter6_end = i
        break

print(f"第六章范围: 段落{chapter6_start} - {chapter6_end}")

# 遍历所有表格，找到第六章中的表格
chapter6_table_count = 0
for table_idx, table in enumerate(doc.tables):
    # 找到表格前一个段落的索引
    prev_para_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p._element is table._element.getprevious():
            prev_para_idx = i
            break
    
    if chapter6_start <= prev_para_idx < chapter6_end:
        chapter6_table_count += 1
        print(f"\n=== 第六章表格{chapter6_table_count} ===")
        
        # 检查表头
        if len(table.rows) > 0:
            for col_idx, cell in enumerate(table.rows[0].cells):
                if cell.paragraphs and cell.paragraphs[0].runs:
                    run = cell.paragraphs[0].runs[0]
                    fs = emu_to_pt(run.font.size)
                    print(f"  表头列{col_idx}: {fs}pt ({get_cn_size(fs)}), 加粗={run.font.bold}, 文本={cell.text[:20]}")
        
        # 检查表体
        if len(table.rows) > 1:
            for row_idx in [1, len(table.rows)//2, len(table.rows)-1]:
                if row_idx < len(table.rows):
                    for col_idx, cell in enumerate(table.rows[row_idx].cells):
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            run = cell.paragraphs[0].runs[0]
                            fs = emu_to_pt(run.font.size)
                            print(f"  第{row_idx}行列{col_idx}: {fs}pt ({get_cn_size(fs)}), 加粗={run.font.bold}, 文本={cell.text[:20]}")
                    break
        
        # 检查表格样式
        print(f"  表格样式: {table.style.name}")
        print(f"  表格行数: {len(table.rows)}")

print(f"\n第六章共{chapter6_table_count}个表格")

# 检查表注样式
print("\n=== 表注样式 ===")
for i, p in enumerate(doc.paragraphs):
    if chapter6_start <= i < chapter6_end and '表6.' in p.text:
        if p.runs:
            run = p.runs[0]
            fs = emu_to_pt(run.font.size)
            print(f"  {p.text}")
            print(f"    字号: {fs}pt ({get_cn_size(fs)}), 加粗={run.font.bold}, 对齐={p.alignment}")
