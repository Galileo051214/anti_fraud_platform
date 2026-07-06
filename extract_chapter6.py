from docx import Document
from docx.oxml.ns import qn
import re

doc_path = r'd:\Project\anti_fraud_platform\docx\《项目阶段二》评审\专业综合课程设计-系统设计说明书-江乐霖-23201317-v2-backup.docx'
doc = Document(doc_path)

def get_paragraph_text(paragraph):
    """获取段落的完整文本"""
    return paragraph.text.strip()

def get_style_name(paragraph):
    """获取段落的样式名称"""
    if paragraph.style:
        return paragraph.style.name
    return 'None'

def is_heading(style_name):
    """判断是否为标题样式"""
    if not style_name:
        return False
    return style_name.startswith('Heading') or '标题' in style_name

def get_heading_level(style_name):
    """获取标题级别"""
    if not style_name:
        return 99
    # 匹配 Heading 1, Heading 2 等
    match = re.match(r'Heading\s*(\d+)', style_name)
    if match:
        return int(match.group(1))
    # 匹配中文标题样式
    if '标题' in style_name:
        match = re.search(r'(\d+)', style_name)
        if match:
            return int(match.group(1))
    return 99

# 首先扫描所有段落，找到第六章的位置
print("=" * 80)
print("扫描文档寻找第六章...")
print("=" * 80)

chapter6_start = -1
chapter6_heading_level = -1

for i, para in enumerate(doc.paragraphs):
    text = get_paragraph_text(para)
    style = get_style_name(para)
    
    # 查找第六章的标题
    if is_heading(style):
        # 匹配 "6" 或 "6.1" 或 "第六章" 开头的标题
        if re.match(r'^6\b', text) or re.match(r'^第[六6]章', text):
            level = get_heading_level(style)
            if chapter6_start == -1:
                chapter6_start = i
                chapter6_heading_level = level
                print(f"找到第六章起始位置: 段落 {i}, 样式: {style}, 级别: {level}, 内容: {text}")
                break

if chapter6_start == -1:
    print("未找到第六章！")
    # 打印所有标题样式的段落，帮助调试
    print("\n文档中的所有标题段落：")
    for i, para in enumerate(doc.paragraphs):
        style = get_style_name(para)
        if is_heading(style):
            print(f"段落 {i}: [{style}] {get_paragraph_text(para)}")
    exit(1)

# 找到第六章的结束位置（下一个同级别或更高级别的标题）
# 改进逻辑：下一个 Heading 1 级别的标题（如7、8、第七章等），但不包括 6.x
chapter6_end = len(doc.paragraphs)
for i in range(chapter6_start + 1, len(doc.paragraphs)):
    para = doc.paragraphs[i]
    style = get_style_name(para)
    text = get_paragraph_text(para)
    
    if is_heading(style):
        level = get_heading_level(style)
        # 如果标题级别小于等于第六章的级别，且不是以 "6." 开头的子章节
        # 或者标题匹配 "7" 或 "第七章" 等新的章节
        if level <= chapter6_heading_level:
            # 检查是否是6.x的子章节
            if not re.match(r'^6\.\d+', text):
                chapter6_end = i
                print(f"找到第六章结束位置: 段落 {i}, 样式: {style}, 级别: {level}, 内容: {text}")
                break

print(f"\n第六章范围: 段落 {chapter6_start} 到 {chapter6_end - 1}, 共 {chapter6_end - chapter6_start} 个段落")
print("=" * 80)

# 提取第六章的段落
print("\n第六章段落内容：")
print("=" * 80)

chapter6_paragraphs = []
for i in range(chapter6_start, chapter6_end):
    para = doc.paragraphs[i]
    text = get_paragraph_text(para)
    style = get_style_name(para)
    level = get_heading_level(style) if is_heading(style) else None
    
    chapter6_paragraphs.append({
        'index': i,
        'text': text,
        'style': style,
        'level': level
    })
    
    level_info = f" [Level: {level}]" if level and level != 99 else ""
    print(f"段落 {i}: [{style}]{level_info}")
    print(f"  内容: {text}")
    print()

# 统计样式分布
print("=" * 80)
print("样式分布统计：")
style_counts = {}
for p in chapter6_paragraphs:
    style = p['style']
    style_counts[style] = style_counts.get(style, 0) + 1

for style, count in sorted(style_counts.items(), key=lambda x: x[1], reverse=True):
    print(f"  {style}: {count} 个段落")

# 提取第六章中的表格
print("\n" + "=" * 80)
print("第六章中的表格：")
print("=" * 80)

# 通过XML元素位置判断表格是否在第六章范围内
body = doc.element.body
body_children = list(body)

start_elem = doc.paragraphs[chapter6_start]._element
end_elem = doc.paragraphs[chapter6_end - 1]._element if chapter6_end < len(doc.paragraphs) else None

try:
    start_idx = body_children.index(start_elem)
    if end_elem:
        end_idx = body_children.index(end_elem)
    else:
        end_idx = len(body_children) - 1
except ValueError:
    start_idx = 0
    end_idx = len(body_children) - 1

tables_in_chapter6 = []
for table_idx, table in enumerate(doc.tables):
    table_element = table._element
    try:
        table_idx_in_body = body_children.index(table_element)
        if start_idx <= table_idx_in_body <= end_idx:
            tables_in_chapter6.append(table)
    except ValueError:
        pass

print(f"\n找到 {len(tables_in_chapter6)} 个表格在第六章中\n")

for table_idx, table in enumerate(tables_in_chapter6):
    print(f"--- 表格 {table_idx + 1} ---")
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  行 {row_idx}: {cells}")
    print()

# 层级关系分析
print("=" * 80)
print("第六章标题层级关系分析：")
print("=" * 80)

for p in chapter6_paragraphs:
    if p['level'] and p['level'] != 99:
        indent = "  " * (p['level'] - 1)
        print(f"{indent}[{p['style']}] {p['text']}")

print("\n" + "=" * 80)
print("分析完成！")
print("=" * 80)
